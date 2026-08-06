# -*- coding: utf-8 -*-
"""Xuất biên bản khảo sát PCCC ra file Word (.docx).

Sử dụng: bien_ban_khao_sat.xuat_bien_ban_khao_sat(data, file_out, nguoi_lap, sdt)
"""
import json
import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


# Mã hóa option → text hiển thị (để phía sales đọc được)
HINH_THUC_LBL = {"so_huu": "Sở hữu", "thue": "Thuê", "khac": "Khác"}
TRANG_THAI_LBL = {"thiet_ke": "Đang thiết kế",
                  "thi_cong": "Đang thi công",
                  "van_hanh": "Đã đi vào vận hành"}
NGUON_NUOC_LBL = {"duong_ong": "Đường ống đô thị",
                  "be_chua": "Bể chứa riêng",
                  "song_ho": "Sông / hồ / ao"}
NGUON_DIEN_LBL = {"1_nguon": "1 nguồn", "2_nguon": "2 nguồn độc lập"}
NGAN_SACH_LBL = {"": "Chưa rõ",
                 "<100tr": "Dưới 100 triệu",
                 "100-500tr": "100 – 500 triệu",
                 "500tr-2ty": "500tr – 2 tỷ",
                 ">2ty": "Trên 2 tỷ"}
TAI_LIEU_LBL = {
    "gp_xay_dung": "Giấy phép xây dựng",
    "gcn_qsdd": "GCN quyền sử dụng đất",
    "bv_thiet_ke_pccc": "Bản vẽ thiết kế PCCC",
    "bv_hoan_cong": "Bản vẽ hoàn công",
    "bb_nghiem_thu_pccc": "Biên bản nghiệm thu PCCC",
    "qd_tham_duyet": "Quyết định thẩm duyệt PCCC trước đó",
    "so_theo_doi_pccc": "Sổ theo dõi PCCC",
    "bb_kiem_tra_ca": "Biên bản kiểm tra của Công an PCCC",
    "hd_su_dung": "Hướng dẫn sử dụng thiết bị PCCC",
    "bh_chay_no": "Hợp đồng bảo hiểm cháy nổ",
}


def _add_p(doc, text="", bold=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def _add_section_header(doc, title):
    """Header section màu trắng nền teal."""
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # set shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "0F766E")
    p._p.get_or_add_pPr().append(shd)


def _kv_table(doc, rows):
    """Tạo bảng 2 cột Key-Value cho danh sách (key, value)."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light List Accent 1"
    for i, (k, v) in enumerate(rows):
        cells = t.rows[i].cells
        cells[0].text = str(k)
        cells[1].text = str(v) if v is not None else ""
        # Bold key
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    # Set column widths
    for row in t.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.5)
    return t


def _bool_lbl(v, true_lbl="✓ Có", false_lbl="☐ Không"):
    return true_lbl if v else false_lbl


def _format_nguoi(ten, chuc_vu, sdt):
    parts = [p.strip() for p in [ten, chuc_vu, sdt] if p and str(p).strip()]
    return " - ".join(parts)


def _format_nguon_nuoc(data):
    """Format nguồn nước multi (từ json) hoặc fallback single."""
    keys = []
    raw_json = data.get("nguon_nuoc_json")
    if raw_json:
        try:
            keys = json.loads(raw_json)
        except Exception:
            keys = []
    if not keys and data.get("nguon_nuoc"):
        keys = [x.strip() for x in str(data["nguon_nuoc"]).split(",")
                if x.strip()]
    return ", ".join(NGUON_NUOC_LBL.get(k, k) for k in keys) or "—"


def xuat_bien_ban_khao_sat(data: dict, file_out: str,
                           nguoi_lap: str = "",
                           sdt_nguoi_lap: str = "") -> str:
    """Tạo biên bản khảo sát Word từ dict data (theo schema sale_lead_khao_sat).
    Trả về đường dẫn file đã ghi."""
    doc = Document()

    # ===== Header công ty =====
    _add_p(doc, "CÔNG TY CỔ PHẦN FIRESMART (HOẶC TÊN CÔNG TY)",
           bold=True, size=12, align="center")
    _add_p(doc, "Địa chỉ: ___ · ĐT: ___ · Website: ___",
           size=10, align="center", color=(0x80, 0x80, 0x80))
    doc.add_paragraph()

    # ===== Tiêu đề =====
    _add_p(doc, "BIÊN BẢN KHẢO SÁT HIỆN TRƯỜNG", bold=True,
           size=16, align="center", color=(0x0F, 0x76, 0x6E))
    _add_p(doc, "(Tư vấn hệ thống PCCC theo QCVN 10:2025/BCA "
           "+ TCVN 7336:2021 + NĐ 105/2025/NĐ-CP)",
           size=10, align="center", color=(0x55, 0x55, 0x55))
    doc.add_paragraph()

    # ===== A. Thông tin chung =====
    _add_section_header(doc, "A. THÔNG TIN CHUNG CÔNG TRÌNH")
    _kv_table(doc, [
        ("Tên công trình", data.get("ten_cong_trinh") or ""),
        ("Địa chỉ", data.get("dia_chi") or ""),
        ("Hình thức sở hữu",
         HINH_THUC_LBL.get(data.get("hinh_thuc_so_huu"), "")),
        ("Trạng thái",
         TRANG_THAI_LBL.get(data.get("trang_thai"), "")),
        ("Tính chất",
         "Công trình độc lập" if data.get("ct_doc_lap")
         else "Một phần của công trình lớn"),
        ("Mã công năng",
         data.get("cong_nang_k") or "(chưa chọn)"),
    ])
    doc.add_paragraph()

    # ===== B. Quy mô + kỹ thuật =====
    _add_section_header(doc, "B. QUY MÔ + KỸ THUẬT CÔNG TRÌNH")
    _kv_table(doc, [
        ("Tổng diện tích sàn",
         f"{data.get('dt_san_tong', 0):.1f} m²"),
        ("Số tầng nổi", str(data.get("so_tang_noi") or 0)),
        ("Số tầng hầm", str(data.get("so_tang_ham") or 0)),
        ("Chiều cao PCCC",
         f"{data.get('cao_pccc', 0):.1f} m"),
        ("Bậc chịu lửa (QCVN 06)",
         data.get("bac_chiu_lua") or ""),
        ("Cấp nguy hiểm cháy KC (S0-S3)",
         data.get("cap_nhc") or ""),
        ("Hạng nguy hiểm cháy (A-E, SX/kho)",
         data.get("hang_nguy_hiem") or ""),
        ("Số người sử dụng dự kiến",
         str(data.get("so_nguoi_du_kien") or 0)),
        ("Khoảng cách tới CT kế bên",
         f"{data.get('kc_ct_ke_ben', 0):.1f} m"),
        ("Đường giao thông xe CC tiếp cận",
         _bool_lbl(data.get("xe_cc_tiep_can"))),
    ])
    doc.add_paragraph()

    # ===== C. Hạ tầng + hệ thống PCCC sẵn có =====
    _add_section_header(doc,
                       "C. HẠ TẦNG KỸ THUẬT + HỆ THỐNG PCCC SẴN CÓ")
    _kv_table(doc, [
        ("Nguồn cấp nước", _format_nguon_nuoc(data)),
        ("Chi tiết nguồn nước",
         data.get("nguon_nuoc_chi_tiet") or ""),
        ("Nguồn điện",
         NGUON_DIEN_LBL.get(data.get("nguon_dien"), "")),
        ("Máy phát điện dự phòng",
         _bool_lbl(data.get("co_may_phat"))),
    ])
    # Bảng hệ thống PCCC sẵn có
    _add_p(doc, "Hệ thống PCCC sẵn có:", bold=True)
    try:
        ht_data = json.loads(data.get("he_thong_sn_json") or "[]")
    except Exception:
        ht_data = []
    if ht_data:
        t = doc.add_table(rows=1 + len(ht_data), cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text = "Hệ thống"
        hdr[1].text = "Tình trạng"
        hdr[2].text = "Hãng / Model"
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for i, ht in enumerate(ht_data):
            r = t.rows[1 + i].cells
            r[0].text = ht.get("ten", "")
            r[1].text = ht.get("tinh_trang", "") or "(chưa khảo sát)"
            r[2].text = ht.get("hang", "")
    doc.add_paragraph()

    # ===== D. Pháp lý + tài liệu =====
    _add_section_header(doc, "D. PHÁP LÝ + TÀI LIỆU KHÁCH HÀNG ĐÃ CÓ")
    try:
        tl_data = json.loads(data.get("tai_lieu_json") or "{}")
    except Exception:
        tl_data = {}
    tl_rows = []
    for k, lbl in TAI_LIEU_LBL.items():
        tl_rows.append((lbl, _bool_lbl(tl_data.get(k))))
    _kv_table(doc, tl_rows)
    _add_p(doc, "Lịch sử PCCC:", bold=True)
    _add_p(doc, data.get("lich_su_pccc") or "(không có ghi nhận)",
           size=11)
    doc.add_paragraph()

    # ===== E. Thương mại + đánh giá =====
    _add_section_header(doc, "E. THƯƠNG MẠI + ĐÁNH GIÁ SALES")
    _kv_table(doc, [
        ("Yêu cầu / mong muốn KH",
         data.get("yc_kh") or ""),
        ("Ngân sách dự kiến",
         NGAN_SACH_LBL.get(data.get("ngan_sach"), "")),
        ("Deadline triển khai",
         data.get("deadline") or ""),
        ("Người ra quyết định",
         _format_nguoi(data.get("qd_ten"), data.get("qd_chuc_vu"),
                       data.get("qd_sdt"))
         or data.get("nguoi_quyet_dinh") or ""),
        ("Người liên hệ kỹ thuật",
         data.get("lien_he_ky_thuat") or ""),
        ("Đối thủ đã báo giá",
         _bool_lbl(data.get("doi_thu_da_bao_gia"))),
        ("Tên đối thủ",
         data.get("doi_thu_ten") or ""),
    ])
    _add_p(doc, "Đánh giá sales:", bold=True)
    _add_p(doc, data.get("danh_gia_sales") or "(chưa đánh giá)",
           size=11)
    _add_p(doc, "Bước tiếp theo:", bold=True)
    _add_p(doc, data.get("buoc_tiep_theo") or "(chưa xác định)",
           size=11)
    doc.add_paragraph()

    # ===== Phần ký =====
    ngay_ks = data.get("ngay_khao_sat") or ""
    if hasattr(ngay_ks, "strftime"):
        ngay_ks = ngay_ks.strftime("%d/%m/%Y")
    elif ngay_ks:
        try:
            d = datetime.datetime.strptime(str(ngay_ks)[:10],
                                          "%Y-%m-%d")
            ngay_ks = d.strftime("%d/%m/%Y")
        except Exception:
            pass
    _add_p(doc, f"Hà Nội, ngày {ngay_ks}", align="right",
           size=11, color=(0x55, 0x55, 0x55))
    doc.add_paragraph()

    # Bảng ký 2 cột
    t = doc.add_table(rows=4, cols=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1 = t.rows[0].cells
    h1[0].text = "ĐẠI DIỆN KHÁCH HÀNG"
    h1[1].text = "ĐẠI DIỆN BÊN TƯ VẤN"
    for cell in h1:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
    h2 = t.rows[1].cells
    h2[0].text = "(Ký, ghi rõ họ tên)"
    h2[1].text = "(Ký, ghi rõ họ tên)"
    for cell in h2:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.italic = True
    # 2 dòng trống để ký
    for _ in range(2):
        rr = t.add_row().cells
        rr[0].text = ""
        rr[1].text = ""
    # Dòng cuối: tên + chức danh
    td_full = _format_nguoi(data.get("td_ten"), data.get("td_chuc_vu"),
                            data.get("td_sdt"))
    nguoi_tiep_don = (td_full or data.get("nguoi_tiep_don")
                      or "___________________")
    nguoi_ks = data.get("nguoi_khao_sat") or nguoi_lap or "___________________"
    last = t.rows[-1].cells
    last[0].text = nguoi_tiep_don
    last[1].text = f"{nguoi_ks}\nSĐT: {sdt_nguoi_lap or '___'}"
    for cell in last:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== Footer =====
    doc.add_paragraph()
    _add_p(doc,
           "Tài liệu này được lập trên phần mềm Tư vấn & Báo giá PCCC — "
           "tham khảo, không thay thế hồ sơ thẩm duyệt PCCC chính thức "
           "của cơ quan có thẩm quyền.",
           size=9, align="center", color=(0x88, 0x88, 0x88))

    doc.save(file_out)
    return file_out
