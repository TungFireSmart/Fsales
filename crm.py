from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QPushButton, QTableWidgetItem, QMessageBox,
    QFileDialog, QTableWidget, QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QLineEdit, QInputDialog, QFormLayout, QHeaderView
)
from PyQt6.QtCore import QDate, Qt, QSize
from PyQt6.QtGui import QPixmap, QIcon

import re
import sys
import json
import subprocess
import webbrowser
import requests
from pathlib import Path
from datetime import datetime
from docx import Document

import misc
import quotation
from UI.ds_cong_ty import Ui_ViewAllCompany
from UI.chi_tiet_cong_ty import Ui_ViewDetailCompany
from UI.supplier_detail import Ui_SupplierDetailWindow
from UI.san_pham_moi import Ui_SanPhamMoi
from UI.supplier_product_editor import Ui_SupplierProductEditor
from ui_theme import apply_ui_v2
import file_handle


class Crm(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        # Không dựng UI tại __init__ để tránh lỗi self.win_companyview chưa tồn tại.
        # Màn CRM được mở qua company_view().
        self.uic8 = None
        self.win_companyview = None
        self.suppliers_tab = None
        self.suppliers_table = None
        self.suppliers_search = None
        self.setWindowTitle(QApplication.translate("CompanyView", "Fsale v1.1.28"))

    def company_view(self):
        self.win_companyview = QMainWindow()
        self.uic8 = Ui_ViewAllCompany()
        self.uic8.setupUi(self.win_companyview)
        apply_ui_v2(self.win_companyview)
        self.win_companyview.show()

        self.uic8.tabWidget.setTabText(self.uic8.tabWidget.indexOf(self.uic8.tab), "Công ty")
        self.uic8.tabWidget.setTabText(self.uic8.tabWidget.indexOf(self.uic8.tab_2), "Cá nhân")
        self.uic8.tabWidget.setTabText(self.uic8.tabWidget.indexOf(self.uic8.tab_3), "Nhà cung cấp")

        self.suppliers_tab = self.uic8.tab_3
        self.suppliers_table = self.uic8.table_suppliers
        self.suppliers_search = self.uic8.txt_supplier_search

        self.suppliers_table.setColumnCount(5)
        self.suppliers_table.setHorizontalHeaderLabels([
            'Mã NCC', 'Tên NCC', 'SĐT', 'Liên hệ', 'Thao tác'
        ])
        self.suppliers_table.verticalHeader().setVisible(False)
        h = self.suppliers_table.horizontalHeader()
        h.setSectionResizeMode(QHeaderView.ResizeMode.Stretch)

        try:
            self.uic8.but_supplier_search.clicked.disconnect()
        except Exception:
            pass
        try:
            self.uic8.but_supplier_refresh.clicked.disconnect()
        except Exception:
            pass
        try:
            self.uic8.but_supplier_add.clicked.disconnect()
        except Exception:
            pass
        try:
            self.suppliers_search.returnPressed.disconnect()
        except Exception:
            pass

        self.uic8.but_supplier_search.clicked.connect(self.search_suppliers)
        self.uic8.but_supplier_refresh.clicked.connect(self.load_suppliers)
        self.uic8.but_supplier_add.clicked.connect(self.add_supplier)
        self.suppliers_search.returnPressed.connect(self.search_suppliers)

        self.uic8.tabWidget.currentChanged.connect(lambda: Crm.on_tab_changed(self, self.uic8.tabWidget.currentIndex()))

        self.uic8.tableWidget.clear()
        self.uic8.tableWidget.verticalHeader().setVisible(False)

        # Tối ưu hiệu năng: chỉ load danh sách gần nhất, tránh kéo toàn bộ bảng
        code = "SELECT * FROM ds_cong_ty ORDER BY ma_kh DESC LIMIT 500"
        result = misc.sql_all(code, None)

        if result:
            self.uic8.tableWidget.setColumnCount(5)
            self.uic8.tableWidget.setRowCount(len(result))
            self.uic8.tableWidget.setHorizontalHeaderLabels(['Mã KH', 'Mã số thuế', 'Tên công ty', 'Địa chỉ', 'Thao tác'])
            self.uic8.tableWidget.setColumnWidth(0, 40)
            self.uic8.tableWidget.setColumnWidth(1, 80)
            self.uic8.tableWidget.setColumnWidth(2, 200)
            self.uic8.tableWidget.setColumnWidth(3, 220)
            self.uic8.tableWidget.setColumnWidth(4, 50)
            for row in range(len(result)):
                self.uic8.tableWidget.setItem(row, 0, QTableWidgetItem(str(result[row][0])))
                self.uic8.tableWidget.setItem(row, 1, QTableWidgetItem(result[row][2]))
                self.uic8.tableWidget.setItem(row, 2, QTableWidgetItem(result[row][1]))
                self.uic8.tableWidget.setItem(row, 3, QTableWidgetItem(result[row][4]))

                but1 = QPushButton('Xem')
                but1.clicked.connect(lambda _, r=row: Crm.view_detail_company(self, result[r][2]))
                self.uic8.tableWidget.setCellWidget(row, 4, but1)

                self.uic8.tableWidget.setRowHeight(row, 70)
        else:
            print('Lỗi không hiển thị được')
            self.uic8.tableWidget.clear()

        # Tránh connect trùng khi mở lại màn hình CRM nhiều lần
        try:
            self.uic8.textEdit.textChanged.disconnect()
        except Exception:
            pass
        try:
            self.uic8.but_search.clicked.disconnect()
        except Exception:
            pass
        try:
            self.uic8.but_add_company.clicked.disconnect()
        except Exception:
            pass

        self.uic8.textEdit.textChanged.connect(lambda: Crm.search_company(self))
        self.uic8.but_search.clicked.connect(lambda: Crm.search_company(self))
        self.uic8.but_add_company.clicked.connect(lambda: Crm.add_company(self))

    @staticmethod
    def normalize_phone(phone_text):
        digits = re.sub(r"\D+", "", (phone_text or ""))
        if digits.startswith("84") and len(digits) in (11, 12):
            digits = "0" + digits[2:]
        if len(digits) == 9 and not digits.startswith("0"):
            digits = "0" + digits
        return digits

    @staticmethod
    def normalize_mst(mst_text):
        txt = (mst_text or "").strip().replace(" ", "")
        txt = txt.replace("–", "-").replace("—", "-")
        return txt

    def _history_data(self, mst='', phone=''):
        mst = Crm.normalize_mst(mst)
        phone = Crm.normalize_phone(phone)

        # 1) Thu thập lead liên quan theo MST/SĐT
        leads = []
        if mst:
            leads = misc.sql_all(
                "SELECT lead_id, phu_trach, sdt FROM sale_lead WHERE mst = %s AND (check_delete IS NULL OR check_delete != '1')",
                (mst,)
            )
        if not leads and phone:
            leads = misc.sql_all(
                "SELECT lead_id, phu_trach, sdt FROM sale_lead WHERE sdt = %s AND (check_delete IS NULL OR check_delete != '1')",
                (phone,)
            )

        lead_ids = sorted({int(r[0]) for r in leads if r and r[0] is not None})

        # 2) Thu thập phone liên quan từ ds_ca_nhan / ds_cong_ty để bắt đủ lịch sử (kể cả record cũ)
        phones = set()
        if phone:
            phones.add(phone)

        if mst:
            cn_rows = misc.sql_all("SELECT dien_thoai FROM ds_ca_nhan WHERE mst_cong_ty = %s", (mst,)) or []
            for rr in cn_rows:
                p = Crm.normalize_phone(rr[0] if rr and len(rr) > 0 else '')
                if p:
                    phones.add(p)

            cty = misc.sql_one("SELECT dien_thoai_cong_ty, sdt_nguoi_lh FROM ds_cong_ty WHERE mst = %s", (mst,))
            if cty:
                for p_raw in cty:
                    p = Crm.normalize_phone(p_raw or '')
                    if p:
                        phones.add(p)

        for rr in leads:
            if len(rr) > 2 and rr[2]:
                p = Crm.normalize_phone(rr[2])
                if p:
                    phones.add(p)

        # 3) Lấy báo giá theo lead_id OR theo phone
        quote_rows = []
        if lead_ids:
            id_csv = ','.join(str(i) for i in lead_ids)
            quote_rows.extend(misc.sql_all(f"SELECT * FROM ds_bao_gia WHERE lead_id IN ({id_csv})", None) or [])

        for p in phones:
            quote_rows.extend(misc.sql_all("SELECT * FROM ds_bao_gia WHERE dien_thoai = %s", (p,)) or [])

        # dedupe báo giá theo so_bg
        quote_map = {}
        for r in quote_rows:
            if r and r[0] is not None:
                quote_map[int(r[0])] = r
        quote_rows = [quote_map[k] for k in sorted(quote_map.keys(), reverse=True)]

        quote_ids = sorted({int(r[0]) for r in quote_rows if r and r[0] is not None})

        # 4) Lấy đơn hàng theo so_bg báo giá
        order_rows = []
        if quote_ids:
            q_csv = ','.join(str(i) for i in quote_ids)
            order_rows = misc.sql_all(f"SELECT * FROM ds_don_hang WHERE so_bg IN ({q_csv}) ORDER BY so_bg DESC", None) or []

        # 5) Xác định đơn đã xuất kho
        delivered_so_bg = set()
        if quote_ids:
            q_csv = ','.join(str(i) for i in quote_ids)
            xk = misc.sql_all(f"SELECT so_bg FROM xuat_kho WHERE so_bg IN ({q_csv}) AND kt_duyet = 'T'", None) or []
            delivered_so_bg = {int(r[0]) for r in xk if r and r[0] is not None}

        # fallback: nếu chưa có bản ghi xuat_kho thì dùng cờ đơn hàng đã hoàn thành/đã giao
        if not delivered_so_bg:
            for r in order_rows:
                try:
                    so = int(r[0])
                    da_giao = str(r[12]).upper() if len(r) > 12 and r[12] is not None else 'F'
                    da_ht = str(r[13]).upper() if len(r) > 13 and r[13] is not None else 'F'
                    if da_giao in ('T', '1') or da_ht in ('T', '1'):
                        delivered_so_bg.add(so)
                except Exception:
                    pass

        order_so_bg = {int(r[0]) for r in order_rows if r and r[0] is not None}
        open_quotes = [r for r in quote_rows if int(r[0]) not in order_so_bg]
        done_quotes = [r for r in quote_rows if (int(r[0]) in order_so_bg) or (str(r[4]).upper() in ('T', '1'))]
        delivered_orders = [r for r in order_rows if int(r[0]) in delivered_so_bg]

        # owner gần nhất: ưu tiên lead.phu_trach, fallback quote.user
        latest_owner = ''
        owners = [r[1] for r in leads if len(r) > 1 and r[1]]
        if owners:
            latest_owner = owners[-1]
        elif quote_rows:
            users = [r[5] for r in quote_rows if len(r) > 5 and r[5]]
            latest_owner = users[0] if users else ''

        total_quotes = len(quote_rows)
        total_orders_delivered = len(delivered_orders)
        close_rate = round((total_orders_delivered / total_quotes) * 100, 2) if total_quotes else 0

        return {
            'latest_owner': latest_owner,
            'quotes': quote_rows,
            'open_quotes': open_quotes,
            'done_quotes': done_quotes,
            'orders': order_rows,
            'delivered_orders': delivered_orders,
            'total_quotes': total_quotes,
            'total_orders_delivered': total_orders_delivered,
            'close_rate': close_rate,
        }

    def _show_quote_popup(self, so_bg):
        kq = misc.sql_one("SELECT so_bg, ngaythang, tieu_de, user, sotien, dat_hang, thanh_toan FROM ds_bao_gia WHERE so_bg = %s", (so_bg,))
        if not kq:
            QMessageBox.information(None, "Báo giá", f"Không tìm thấy báo giá #{so_bg}")
            return
        txt = (
            f"Số BG: {kq[0]}\n"
            f"Ngày: {kq[1]}\n"
            f"Tiêu đề: {kq[2]}\n"
            f"Người làm BG: {kq[3]}\n"
            f"Tổng tiền: {kq[4]:,} VNĐ\n"
            f"Đặt hàng: {kq[5]} | Thanh toán: {kq[6]}"
        )
        QMessageBox.information(None, f"Báo giá #{so_bg}", txt)

    def _show_order_popup(self, so_bg):
        kq = misc.sql_one("SELECT so_bg, lead_id, tien_hang, vat, da_thanh_toan, da_giao_hang, da_hoan_thanh, nguoi_tu_van, nguoi_cai_dat FROM ds_don_hang WHERE so_bg = %s", (so_bg,))
        if not kq:
            QMessageBox.information(None, "Đơn hàng", f"Không tìm thấy đơn hàng theo BG #{so_bg}")
            return
        txt = (
            f"Số BG: {kq[0]}\n"
            f"Lead ID: {kq[1]}\n"
            f"Tiền hàng: {kq[2]:,} VNĐ | VAT: {kq[3]:,} VNĐ\n"
            f"Đã thanh toán: {kq[4]} | Đã giao hàng: {kq[5]} | Hoàn thành: {kq[6]}\n"
            f"Tư vấn: {kq[7]} | Cài đặt: {kq[8]}"
        )
        QMessageBox.information(None, f"Đơn hàng BG #{so_bg}", txt)

    def _render_history_tab(self, uic, mst='', phone=''):
        data = Crm._history_data(self, mst=mst, phone=phone)

        summary = (
            f"Tóm tắt: {data['total_quotes']} báo giá | "
            f"{data['total_orders_delivered']} đơn đã xuất kho | "
            f"Tỷ lệ chốt: {data['close_rate']}% | "
            f"Phụ trách gần nhất: {data['latest_owner'] or 'Chưa rõ'}"
        )
        uic.label_noti.setText(summary)

        rows = []
        # 1) Báo giá chưa chốt
        for r in data['open_quotes']:
            rows.append(['Báo giá chưa chốt', str(r[0]), str(r[2]), (r[6] or '')[:80], str(r[3] or ''), 'quote'])

        # 2) Đơn hàng đã thực hiện (đã xuất kho)
        for r in data['delivered_orders']:
            rows.append(['Đơn đã xuất kho', str(r[0]), '', f"Tiền hàng: {r[2]:,}", str(r[16] or ''), 'order'])

        # 3) Báo giá đã thực hiện
        for r in data['done_quotes']:
            rows.append(['Báo giá đã thực hiện', str(r[0]), str(r[2]), (r[6] or '')[:80], str(r[3] or ''), 'quote'])

        uic.tableWidget.clear()
        uic.tableWidget.setColumnCount(6)
        uic.tableWidget.setRowCount(len(rows))
        uic.tableWidget.setHorizontalHeaderLabels(['Loại', 'Mã', 'Ngày', 'Nội dung', 'Người phụ trách', 'Thao tác'])
        uic.tableWidget.setColumnWidth(0, 140)
        uic.tableWidget.setColumnWidth(1, 70)
        uic.tableWidget.setColumnWidth(2, 90)
        uic.tableWidget.setColumnWidth(3, 220)
        uic.tableWidget.setColumnWidth(4, 120)
        uic.tableWidget.setColumnWidth(5, 80)

        for i, rr in enumerate(rows):
            uic.tableWidget.setItem(i, 0, QTableWidgetItem(rr[0]))
            uic.tableWidget.setItem(i, 1, QTableWidgetItem(rr[1]))
            uic.tableWidget.setItem(i, 2, QTableWidgetItem(rr[2]))
            uic.tableWidget.setItem(i, 3, QTableWidgetItem(rr[3]))
            uic.tableWidget.setItem(i, 4, QTableWidgetItem(rr[4]))

            but = QPushButton('Xem')
            if rr[5] == 'order':
                but.clicked.connect(lambda _, so_bg=rr[1]: Crm._show_order_popup(self, int(so_bg)))
            else:
                but.clicked.connect(lambda _, so_bg=rr[1]: Crm._show_quote_popup(self, int(so_bg)))
            uic.tableWidget.setCellWidget(i, 5, but)
            uic.tableWidget.setRowHeight(i, 58)

        uic.tableWidget.repaint()

    def load_suppliers(self, keyword=''):
        try:
            if keyword:
                q = (
                    "SELECT supplier_id, supplier_code, name, phone, contact_name, supplier_group, status "
                    "FROM fs_suppliers WHERE name LIKE %s OR supplier_code LIKE %s OR phone LIKE %s OR tax_code LIKE %s "
                    "ORDER BY supplier_id DESC LIMIT 500"
                )
                k = f"%{keyword}%"
                rows = misc.sql_all(q, (k, k, k, k))
            else:
                rows = misc.sql_all(
                    "SELECT supplier_id, supplier_code, name, phone, contact_name, supplier_group, status "
                    "FROM fs_suppliers ORDER BY supplier_id DESC LIMIT 500",
                    None,
                )
        except Exception as e:
            QMessageBox.warning(None, "Nhà cung cấp", f"Chưa có bảng NCC hoặc lỗi truy vấn: {e}")
            return

        rows = rows or []
        self.suppliers_table.setRowCount(len(rows))
        for r, item in enumerate(rows):
            sid = int(item[0])
            self.suppliers_table.setItem(r, 0, QTableWidgetItem(str(item[1] or '')))
            self.suppliers_table.setItem(r, 1, QTableWidgetItem(str(item[2] or '')))
            self.suppliers_table.setItem(r, 2, QTableWidgetItem(str(item[3] or '')))
            self.suppliers_table.setItem(r, 3, QTableWidgetItem(str(item[4] or '')))

            but_view = QPushButton("Xem")
            but_view.clicked.connect(lambda _, supplier_id=sid: self.quick_edit_supplier(supplier_id))
            self.suppliers_table.setCellWidget(r, 4, but_view)

            self.suppliers_table.setRowHeight(r, 56)

    def search_suppliers(self):
        key = (self.suppliers_search.text() or '').strip()
        self.load_suppliers(key)

    def _sync_supplier_to_crm(self, name, tax_code='', contact_name='', phone='', email='', address=''):
        tax_code = Crm.normalize_mst(tax_code)
        phone = Crm.normalize_phone(phone)

        if tax_code:
            # upsert company master theo MST
            kq = misc.sql_one("SELECT ma_kh FROM ds_cong_ty WHERE mst = %s", (tax_code,))
            if kq:
                misc.sql_commit(
                    "UPDATE ds_cong_ty SET ten_cong_ty=%s, dia_chi=%s, dien_thoai_cong_ty=%s, email_cong_ty=%s, nguoi_lien_he=%s, sdt_nguoi_lh=%s WHERE mst=%s",
                    (name, address, phone, email, contact_name, phone, tax_code)
                )
            else:
                misc.sql_commit(
                    "INSERT INTO ds_cong_ty (ten_cong_ty, mst, dia_chi, dien_thoai_cong_ty, email_cong_ty, nguoi_lien_he, sdt_nguoi_lh) VALUES (%s,%s,%s,%s,%s,%s,%s)",
                    (name, tax_code, address, phone, email, contact_name, phone)
                )

        if phone and contact_name:
            # upsert contact theo SĐT
            kq_cn = misc.sql_one("SELECT dien_thoai FROM ds_ca_nhan WHERE dien_thoai = %s", (phone,))
            if kq_cn:
                misc.sql_commit(
                    "UPDATE ds_ca_nhan SET ten=%s, email=%s, ten_cong_ty=%s, mst_cong_ty=%s, address=%s WHERE dien_thoai=%s",
                    (contact_name, email, name, tax_code, address, phone)
                )
            else:
                misc.sql_commit(
                    "INSERT INTO ds_ca_nhan (ten, dien_thoai, email, ten_cong_ty, mst_cong_ty, address) VALUES (%s,%s,%s,%s,%s,%s)",
                    (contact_name, phone, email, name, tax_code, address)
                )

    def add_supplier(self):
        name, ok = QInputDialog.getText(self.win_companyview, "Thêm NCC", "Tên nhà cung cấp:")
        if not ok or not str(name).strip():
            return

        contact, ok = QInputDialog.getText(self.win_companyview, "Thêm NCC", "Người liên hệ (bắt buộc):")
        if not ok or not str(contact).strip():
            QMessageBox.warning(self.win_companyview, "Thêm NCC", "Phải có người liên hệ")
            return

        phone, _ = QInputDialog.getText(self.win_companyview, "Thêm NCC", "Số điện thoại liên hệ:")
        email, _ = QInputDialog.getText(self.win_companyview, "Thêm NCC", "Email liên hệ:")
        tax_code, _ = QInputDialog.getText(self.win_companyview, "Thêm NCC", "Mã số thuế:")
        address, _ = QInputDialog.getText(self.win_companyview, "Thêm NCC", "Địa chỉ:")

        if not str(phone).strip() and not str(email).strip():
            QMessageBox.warning(self.win_companyview, "Thêm NCC", "Phải có ít nhất SĐT hoặc email liên hệ")
            return

        try:
            max_code = misc.sql_one("SELECT supplier_code FROM fs_suppliers WHERE supplier_code LIKE 'NCC-%' ORDER BY supplier_id DESC LIMIT 1", None)
            if not max_code or not max_code[0]:
                code = 'NCC-0001'
            else:
                try:
                    n = int(str(max_code[0]).split('-')[-1]) + 1
                except Exception:
                    n = 1
                code = f"NCC-{n:04d}"

            phone_norm = Crm.normalize_phone(phone)
            misc.sql_commit(
                "INSERT INTO fs_suppliers (supplier_code, name, tax_code, phone, email, contact_name, address, supplier_group, status) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                (code, name.strip(), Crm.normalize_mst(tax_code), phone_norm, str(email).strip(), contact.strip(), str(address).strip(), 'Khác', 'Đang hợp tác')
            )
            self._sync_supplier_to_crm(name.strip(), tax_code, contact.strip(), phone_norm, str(email).strip(), str(address).strip())

            self.load_suppliers()
            self.uic8.label_noti.setText(f"Đã thêm NCC: {code} - {name.strip()} (đã sync CRM)")
        except Exception as e:
            QMessageBox.warning(self.win_companyview, "Thêm NCC", f"Không thể thêm NCC: {e}")

    def quick_edit_supplier(self, supplier_id: int):
        row = misc.sql_one(
            "SELECT supplier_code, name, phone, contact_name, supplier_group, status, tax_code, email, address, province, notes FROM fs_suppliers WHERE supplier_id=%s",
            (supplier_id,),
        )
        if not row:
            QMessageBox.warning(self.win_companyview, "Nhà cung cấp", "Không tìm thấy NCC")
            return

        win = QMainWindow(self.win_companyview)
        ui = Ui_SupplierDetailWindow()
        ui.setupUi(win)
        apply_ui_v2(win)
        win.setWindowTitle(f"Chi tiết NCC {row[0]} - {row[1]}")

        ui.txt_name.setText(str(row[1] or ''))
        ui.txt_contact.setText(str(row[3] or ''))
        ui.txt_phone.setText(str(row[2] or ''))
        ui.txt_email.setText(str(row[7] or ''))
        ui.txt_tax.setText(str(row[6] or ''))
        ui.txt_address.setText(str(row[8] or ''))
        ui.txt_province.setText(str(row[9] or ''))
        ui.txt_group.setText(str(row[4] or 'Khác'))
        ui.txt_status.setText(str(row[5] or 'Đang hợp tác'))
        ui.txt_notes.setText(str(row[10] or ''))

        # soft-delete cột ẩn sản phẩm (không xóa vật lý)
        col_hidden = misc.sql_one(
            "SELECT COUNT(*) FROM information_schema.COLUMNS WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME='fs_supplier_products' AND COLUMN_NAME='is_hidden'",
            None,
        )
        if not col_hidden or int(col_hidden[0] or 0) == 0:
            misc.sql_commit("ALTER TABLE fs_supplier_products ADD COLUMN is_hidden TINYINT(1) NOT NULL DEFAULT 0", None)

        ui.table_products.setColumnCount(7)
        ui.table_products.setHorizontalHeaderLabels(["Mã hàng", "Tên hàng", "Giá nhập", "Nhãn hiệu", "NSX", "Xuất xứ", "Thao tác"])
        ui.table_products.setColumnWidth(0, 90)
        ui.table_products.setColumnWidth(1, 220)
        ui.table_products.setColumnWidth(2, 100)
        ui.table_products.setColumnWidth(3, 110)
        ui.table_products.setColumnWidth(4, 110)
        ui.table_products.setColumnWidth(5, 110)
        ui.table_products.setColumnWidth(6, 60)

        # File báo giá NCC đã tách sang màn hình riêng

        def _hide_item(item_code: str):
            if not item_code:
                return
            ok = QMessageBox.question(
                win,
                "Xóa SP",
                f"<span style='color:red; font-weight:600;'>Chắc chắn muốn xóa mã hàng {item_code} ?</span>",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No,
            )
            if ok != QMessageBox.StandardButton.Yes:
                return
            misc.sql_commit(
                "UPDATE fs_supplier_products SET is_hidden=1, updated_at=NOW() WHERE supplier_id=%s AND item_code=%s",
                (supplier_id, item_code),
            )
            load_items()

        def load_items():
            rows = misc.sql_all(
                "SELECT item_code, item_name, model, brand, manufacturer, origin, latest_price FROM fs_supplier_products WHERE supplier_id=%s AND IFNULL(is_hidden,0)=0 ORDER BY id DESC",
                (supplier_id,),
            ) or []
            ui.table_products.setRowCount(len(rows))
            for i, it in enumerate(rows):
                item_code = str(it[0] or '')
                ui.table_products.setItem(i, 0, QTableWidgetItem(item_code))
                ui.table_products.setItem(i, 1, QTableWidgetItem(str(it[1] or '')))
                ui.table_products.setItem(i, 2, QTableWidgetItem(f"{int(it[6] or 0):,}"))
                ui.table_products.setItem(i, 3, QTableWidgetItem(str(it[3] or '')))
                ui.table_products.setItem(i, 4, QTableWidgetItem(str(it[4] or '')))
                ui.table_products.setItem(i, 5, QTableWidgetItem(str(it[5] or '')))
                but_del = QPushButton("Xóa SP")
                but_del.clicked.connect(lambda _, code=item_code: _hide_item(code))
                ui.table_products.setCellWidget(i, 6, but_del)
                ui.table_products.setRowHeight(i, 40)

        def add_item():
            self.win_supplier_product = QMainWindow(win)
            self.ui_supplier_product = Ui_SupplierProductEditor()
            self.ui_supplier_product.setupUi(self.win_supplier_product)
            apply_ui_v2(self.win_supplier_product)
            self.win_supplier_product.setWindowTitle(f"Thêm/Cập nhật mã hàng NCC - {row[0]}")

            u = self.ui_supplier_product
            u.but_save.setText("Lưu mã hàng cho NCC")
            u.but_update.hide()
            u.txt_search.setPlaceholderText("Nhập model (1 dòng) rồi Enter để tìm")
            u.list_images.setViewMode(u.list_images.ViewMode.IconMode)
            u.list_images.setIconSize(QSize(72, 72))
            u.list_images.setResizeMode(u.list_images.ResizeMode.Adjust)
            u.list_images.setGridSize(QSize(90, 110))
            u.list_images.setSpacing(8)

            # Bỏ cụm giá cho thuê khỏi màn hình NCC
            if hasattr(u, 'groupBox') and u.groupBox is not None:
                u.groupBox.hide()

            # Bổ sung trường bắt buộc ngay trên màn hình
            supplier_row = misc.sql_one("SELECT name, phone, contact_name FROM fs_suppliers WHERE supplier_id=%s", (supplier_id,))
            if isinstance(supplier_row, dict):
                _sup_name = str(supplier_row.get('name') or '')
                _sup_phone = str(supplier_row.get('phone') or '')
                _sup_contact = str(supplier_row.get('contact_name') or '')
            elif isinstance(supplier_row, (list, tuple)):
                _sup_name = str(supplier_row[0] or '') if len(supplier_row) > 0 else ''
                _sup_phone = str(supplier_row[1] or '') if len(supplier_row) > 1 else ''
                _sup_contact = str(supplier_row[2] or '') if len(supplier_row) > 2 else ''
            else:
                _sup_name = ''
                _sup_phone = ''
                _sup_contact = ''

            u.txt_supplier_name.setText(_sup_name)
            u.txt_supplier_phone.setText(_sup_phone)
            u.txt_supplier_contact.setText(_sup_contact)

            misc.sql_commit(
                """
                CREATE TABLE IF NOT EXISTS fs_supplier_product_images (
                    image_id BIGINT NOT NULL AUTO_INCREMENT,
                    supplier_id BIGINT NOT NULL,
                    item_code VARCHAR(128) NOT NULL,
                    file_name VARCHAR(255) NOT NULL,
                    drive_file_id VARCHAR(128) NOT NULL,
                    drive_link TEXT NOT NULL,
                    created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    is_deleted TINYINT(1) NOT NULL DEFAULT 0,
                    PRIMARY KEY (image_id),
                    KEY idx_spi_supplier_item (supplier_id, item_code)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
                """,
                None,
            )

            def _load_images(item_code: str):
                u.list_images.clear()
                u.lbl_image_preview.setText('(Không có ảnh)')
                u.lbl_image_preview.setPixmap(QPixmap())
                if not item_code:
                    return []
                rows_img = misc.sql_all(
                    "SELECT image_id, file_name, drive_file_id, drive_link FROM fs_supplier_product_images WHERE supplier_id=%s AND item_code=%s AND IFNULL(is_deleted,0)=0 ORDER BY image_id DESC",
                    (supplier_id, item_code),
                ) or []
                from PyQt6.QtWidgets import QListWidgetItem
                for rr in rows_img:
                    image_id = int(rr[0])
                    file_name = str(rr[1] or '')
                    drive_file_id = str(rr[2] or '')
                    drive_link = str(rr[3] or '')
                    li = QListWidgetItem(file_name)
                    li.setData(Qt.ItemDataRole.UserRole, (image_id, drive_file_id, drive_link))

                    # thumbnail icon
                    try:
                        url = f"https://drive.google.com/uc?export=download&id={drive_file_id}"
                        resp = requests.get(url, timeout=15)
                        resp.raise_for_status()
                        pm = QPixmap()
                        pm.loadFromData(resp.content)
                        if not pm.isNull():
                            li.setIcon(QIcon(pm.scaled(72, 72, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)))
                    except Exception:
                        pass

                    u.list_images.addItem(li)
                if u.list_images.count() > 0:
                    u.list_images.setCurrentRow(0)
                return rows_img

            def _show_selected_image():
                cur = u.list_images.currentItem()
                if not cur:
                    u.lbl_image_preview.setText('(Không có ảnh)')
                    u.lbl_image_preview.setPixmap(QPixmap())
                    return
                _img_id, drive_file_id, _drive_link = cur.data(Qt.ItemDataRole.UserRole)
                try:
                    url = f"https://drive.google.com/uc?export=download&id={drive_file_id}"
                    resp = requests.get(url, timeout=20)
                    resp.raise_for_status()
                    pm = QPixmap()
                    pm.loadFromData(resp.content)
                    if pm.isNull():
                        raise RuntimeError('Ảnh không hợp lệ')
                    u.lbl_image_preview.setPixmap(pm.scaled(u.lbl_image_preview.width()-4, u.lbl_image_preview.height()-4, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
                except Exception:
                    u.lbl_image_preview.setText('Không tải được ảnh')
                    u.lbl_image_preview.setPixmap(QPixmap())

            def _upload_new_image_for_current_product(replace_selected=False):
                item_code = u.txt_model.toPlainText().strip()
                item_name = u.txt_ten_sp.toPlainText().strip()
                if not item_code or not item_name:
                    u.label_noti.setStyleSheet('color: red')
                    u.label_noti.setText('Cần nhập model và tên hàng trước khi upload ảnh')
                    return

                uploaded = file_handle.upload_file()  # tenfile|file_id|mime
                if not uploaded:
                    return
                parts = str(uploaded).split('|')
                if len(parts) < 2:
                    u.label_noti.setStyleSheet('color: red')
                    u.label_noti.setText('Upload ảnh thất bại')
                    return
                old_name, file_id = parts[0], parts[1]

                # đổi tên file theo "model + tên hàng"
                desired_name = f"{item_code} {item_name}".strip()
                try:
                    service = file_handle.authenticate_drive_with_client_info()
                    service.files().update(fileId=file_id, body={'name': desired_name}).execute()
                    file_name = desired_name
                except Exception:
                    file_name = old_name

                drive_link = f"https://drive.google.com/file/d/{file_id}/view"

                if replace_selected:
                    cur = u.list_images.currentItem()
                    if cur:
                        old_image_id, old_drive_id, _ = cur.data(Qt.ItemDataRole.UserRole)
                        misc.sql_commit("UPDATE fs_supplier_product_images SET is_deleted=1 WHERE image_id=%s", (old_image_id,))
                        try:
                            file_handle.delete_file_from_drive(old_drive_id)
                        except Exception:
                            pass

                misc.sql_commit(
                    "INSERT INTO fs_supplier_product_images (supplier_id, item_code, file_name, drive_file_id, drive_link, is_deleted) VALUES (%s,%s,%s,%s,%s,0)",
                    (supplier_id, item_code, file_name, file_id, drive_link),
                )
                _load_images(item_code)
                _show_selected_image()

            def _delete_selected_image():
                cur = u.list_images.currentItem()
                if not cur:
                    return
                image_id, drive_id, _ = cur.data(Qt.ItemDataRole.UserRole)
                misc.sql_commit("UPDATE fs_supplier_product_images SET is_deleted=1 WHERE image_id=%s", (image_id,))
                try:
                    file_handle.delete_file_from_drive(drive_id)
                except Exception:
                    pass
                _load_images(u.txt_model.toPlainText().strip())
                _show_selected_image()

            def _save_item_from_form():
                item_code = u.txt_model.toPlainText().strip()
                item_name = u.txt_ten_sp.toPlainText().strip()
                latest_price = float(re.sub(r"\D", "", u.txt_gia_von.toPlainText()) or 0)
                model = item_code  # mã hàng được tính là model
                brand = u.txt_nhan_hieu.toPlainText().strip()
                manufacturer = u.txt_manufacturer.text().strip()
                origin = u.txt_xuat_xu.toPlainText().strip()
                supplier_name = u.txt_supplier_name.text().strip()
                supplier_phone = u.txt_supplier_phone.text().strip()
                supplier_contact = u.txt_supplier_contact.text().strip()

                if not item_code:
                    u.label_noti.setStyleSheet('color: red')
                    u.label_noti.setText('Thiếu model/mã hàng')
                    return
                if latest_price <= 0:
                    u.label_noti.setStyleSheet('color: red')
                    u.label_noti.setText('Thiếu giá hợp lệ')
                    return
                if not manufacturer:
                    u.label_noti.setStyleSheet('color: red')
                    u.label_noti.setText('Thiếu tên nhà sản xuất')
                    return
                if not supplier_name or not supplier_phone or not supplier_contact:
                    u.label_noti.setStyleSheet('color: red')
                    u.label_noti.setText('Thiếu thông tin NCC (tên/SĐT/người liên hệ)')
                    return

                # đồng bộ thông tin NCC bắt buộc
                misc.sql_commit(
                    "UPDATE fs_suppliers SET name=%s, phone=%s, contact_name=%s, updated_at=NOW() WHERE supplier_id=%s",
                    (supplier_name, supplier_phone, supplier_contact, supplier_id),
                )

                exists = misc.sql_one(
                    "SELECT id FROM fs_supplier_products WHERE supplier_id=%s AND item_code=%s",
                    (supplier_id, item_code),
                )
                if exists:
                    misc.sql_commit(
                        "UPDATE fs_supplier_products SET item_name=%s, model=%s, brand=%s, manufacturer=%s, origin=%s, latest_price=%s, is_hidden=0, updated_at=NOW() WHERE supplier_id=%s AND item_code=%s",
                        (item_name, model, brand, manufacturer, origin, latest_price, supplier_id, item_code),
                    )
                else:
                    misc.sql_commit(
                        "INSERT INTO fs_supplier_products (supplier_id, item_code, item_name, model, brand, manufacturer, origin, latest_price) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
                        (supplier_id, item_code, item_name, model, brand, manufacturer, origin, latest_price),
                    )

                u.label_noti.setStyleSheet('color: blue')
                u.label_noti.setText('Đã lưu mã hàng + thông tin NCC')
                load_items()
                _load_images(item_code)
                _show_selected_image()

            def _search_prefill():
                raw = u.txt_search.toPlainText()
                txt = (raw.splitlines()[0] if raw else '').strip()
                if raw != txt:
                    u.txt_search.setText(txt)
                if not txt:
                    return
                # ưu tiên từ fs_supplier_products hiện có của NCC
                kq = misc.sql_one(
                    "SELECT item_code, item_name, latest_price, brand, manufacturer, origin FROM fs_supplier_products WHERE supplier_id=%s AND item_code=%s AND IFNULL(is_hidden,0)=0",
                    (supplier_id, txt),
                )
                if kq:
                    u.txt_model.setText(str(kq[0] or ''))
                    u.txt_ten_sp.setText(str(kq[1] or ''))
                    u.txt_gia_von.setText(f"{int(kq[2] or 0):,}")
                    u.txt_nhan_hieu.setText(str(kq[3] or ''))
                    u.txt_manufacturer.setText(str(kq[4] or ''))
                    u.txt_xuat_xu.setText(str(kq[5] or ''))
                    _load_images(str(kq[0] or ''))
                    _show_selected_image()
                    u.label_noti.setText('Đã nạp dữ liệu mã hàng của NCC')
                    return
                # fallback từ master sản phẩm
                m = misc.sql_one("SELECT ten_san_pham, model, gia_dau_vao, nhan_hieu, xuat_xu FROM gia_tong_hop WHERE model=%s", (txt,))
                if m:
                    u.txt_ten_sp.setText(str(m[0] or ''))
                    u.txt_model.setText(str(m[1] or ''))
                    u.txt_gia_von.setText(f"{int(m[2] or 0):,}")
                    u.txt_nhan_hieu.setText(str(m[3] or ''))
                    u.txt_xuat_xu.setText(str(m[4] or ''))
                    _load_images(str(m[1] or txt))
                    _show_selected_image()
                    u.label_noti.setText('Đã nạp từ danh mục sản phẩm chung')
                else:
                    u.label_noti.setStyleSheet('color: red')
                    u.label_noti.setText('Không tìm thấy model')

            u.but_save.clicked.connect(_save_item_from_form)
            u.but_search.clicked.connect(_search_prefill)
            u.but_add_image.clicked.connect(lambda: _upload_new_image_for_current_product(False))
            u.but_change_image.clicked.connect(lambda: _upload_new_image_for_current_product(True))
            u.but_delete_image.clicked.connect(_delete_selected_image)
            u.list_images.currentItemChanged.connect(lambda *_: _show_selected_image())

            # Search 1 dòng: Enter = tìm, không xuống dòng
            _orig_keypress = u.txt_search.keyPressEvent
            def _search_keypress(ev):
                if ev.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                    _search_prefill()
                    ev.accept()
                    return
                _orig_keypress(ev)
            u.txt_search.keyPressEvent = _search_keypress

            self.win_supplier_product.show()

        def save_supplier():
            name = ui.txt_name.text().strip()
            contact = ui.txt_contact.text().strip()
            phone = Crm.normalize_phone(ui.txt_phone.text())
            email = ui.txt_email.text().strip()
            tax = Crm.normalize_mst(ui.txt_tax.text())
            address = ui.txt_address.text().strip()

            if not name or not contact or (not phone and not email):
                QMessageBox.warning(win, "Nhà cung cấp", "Cần Tên NCC + Người liên hệ + (SĐT hoặc Email)")
                return

            misc.sql_commit(
                "UPDATE fs_suppliers SET name=%s, phone=%s, contact_name=%s, tax_code=%s, email=%s, address=%s, province=%s, supplier_group=%s, status=%s, notes=%s WHERE supplier_id=%s",
                (
                    name,
                    phone,
                    contact,
                    tax,
                    email,
                    address,
                    ui.txt_province.text().strip(),
                    ui.txt_group.text().strip() or 'Khác',
                    ui.txt_status.text().strip() or 'Đang hợp tác',
                    ui.txt_notes.text().strip(),
                    supplier_id,
                ),
            )
            self._sync_supplier_to_crm(name, tax, contact, phone, email, address)
            self.load_suppliers()
            QMessageBox.information(win, "Nhà cung cấp", "Đã lưu thông tin NCC")

        ui.but_add_item.clicked.connect(add_item)
        ui.but_reload_item.clicked.connect(load_items)
        ui.but_save.clicked.connect(save_supplier)

        ui.but_upload_quote.setText("Báo giá NCC")
        ui.but_upload_quote.clicked.connect(lambda: self.open_supplier_quote_files(supplier_id, win))
        for _legacy_name in (
            'but_reload_files',
            'but_open_file_link',
            'but_mark_done',
            'table_files',
            'label_files',
        ):
            _w = getattr(ui, _legacy_name, None)
            if _w is not None:
                _w.hide()

        load_items()
        win.show()
        self.win_supplier_detail = win

    def open_supplier_quote_files(self, supplier_id: int, parent=None):
        supplier = misc.sql_one("SELECT supplier_code, name FROM fs_suppliers WHERE supplier_id=%s", (supplier_id,))
        if not supplier:
            QMessageBox.warning(self.win_companyview, "Báo giá NCC", "Không tìm thấy NCC")
            return

        win = QMainWindow(parent or self.win_companyview)
        win.setWindowTitle(f"Báo giá NCC - {supplier[0]} - {supplier[1]}")
        root = QWidget(win)
        lay = QVBoxLayout(root)

        tb = QTableWidget()
        tb.setColumnCount(7)
        tb.setHorizontalHeaderLabels(["ID", "Thời gian", "Tên file", "Ghi chú", "Trạng thái", "Nhập DB", "Download"])
        tb.setColumnWidth(0, 55)
        tb.setColumnWidth(1, 130)
        tb.setColumnWidth(2, 300)
        tb.setColumnWidth(3, 220)
        tb.setColumnWidth(4, 90)
        tb.setColumnWidth(5, 110)
        tb.setColumnWidth(6, 110)
        lay.addWidget(tb)

        def _extract_drive_id(link: str) -> str:
            if not link:
                return ''
            m = re.search(r'/d/([^/]+)/', link)
            if m:
                return m.group(1)
            m = re.search(r'id=([^&]+)', link)
            return m.group(1) if m else ''

        def _download_one(file_name: str, link: str):
            drive_id = _extract_drive_id(link)
            if not drive_id:
                QMessageBox.warning(win, "Download", "Không đọc được Drive file id")
                return

            folder = QFileDialog.getExistingDirectory(win, "Chọn thư mục lưu file")
            if not folder:
                return

            safe_name = file_name or f"{drive_id}.bin"
            save_path = Path(folder) / safe_name
            url = f"https://drive.google.com/uc?export=download&id={drive_id}"
            try:
                r = requests.get(url, timeout=90)
                r.raise_for_status()
                save_path.write_bytes(r.content)
                QMessageBox.information(win, "Download", f"Đã tải file về:\n{save_path}")
            except Exception as e:
                QMessageBox.warning(win, "Download", f"Tải file thất bại: {e}")

        def _import_one(file_id: int, link: str):
            drive_id = _extract_drive_id(link)
            if not drive_id:
                QMessageBox.warning(win, "Nhập DB", "Không đọc được Drive file id")
                return
            py = r"D:\Fsales_PCCC\.venv\Scripts\python.exe"
            cli = r"C:\Users\Admin\.openclaw\workspace\fsales_connector\cli.py"
            cmd = [py, cli, "import-supplier-quote", "--drive-file-id", drive_id]
            try:
                r = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
                if r.returncode != 0:
                    raise RuntimeError(r.stderr or r.stdout or 'Import thất bại')

                result = {}
                try:
                    result = json.loads((r.stdout or '').strip() or '{}')
                except Exception:
                    result = {}

                _load()
                missing_count = int(result.get('missing_meta_count') or 0)
                process_status = str(result.get('process_status') or '').strip()
                if process_status == 'verify_pending' or missing_count > 0:
                    QMessageBox.warning(
                        win,
                        "Cần xác minh",
                        f"Đã nhập báo giá nhưng còn {missing_count} dòng thiếu model/nhãn hiệu/NSX/xuất xứ.\n"
                        "Vui lòng bổ sung trước khi chốt dữ liệu."
                    )
                else:
                    QMessageBox.information(win, "Nhập DB", "Đã nhập báo giá vào DB thành công")
            except Exception as e:
                QMessageBox.warning(win, "Nhập DB", f"Không nhập được DB: {e}")

        def _open_verify(file_id: int):
            rows = misc.sql_all(
                """
                SELECT
                    l.supplier_quote_line_id,
                    l.line_no,
                    l.item_code,
                    l.item_name,
                    IFNULL(l.model,''),
                    IFNULL(l.brand,''),
                    IFNULL(l.manufacturer,''),
                    IFNULL(l.origin,'')
                FROM fs_supplier_quotes q
                JOIN fs_supplier_quote_lines l ON l.supplier_quote_id=q.supplier_quote_id
                WHERE q.source_file_id=%s
                  AND (
                    IFNULL(l.model,'')='' OR IFNULL(l.brand,'')='' OR
                    IFNULL(l.manufacturer,'')='' OR IFNULL(l.origin,'')=''
                  )
                ORDER BY l.line_no ASC
                """,
                (file_id,),
            ) or []

            if not rows:
                misc.sql_commit(
                    "UPDATE fs_supplier_files SET process_status='done', processed_by='Anna', processed_at=NOW() WHERE file_id=%s",
                    (file_id,),
                )
                _load()
                QMessageBox.information(win, "Xác minh", "Không còn dòng thiếu thông tin. Đã chuyển trạng thái hoàn tất.")
                return

            vwin = QMainWindow(win)
            vwin.setWindowTitle(f"Bổ sung thông tin sản phẩm thiếu - file #{file_id}")
            root2 = QWidget(vwin)
            lay2 = QVBoxLayout(root2)

            tip = QLabel("Các ô màu vàng là bắt buộc: model, nhãn hiệu, nhà sản xuất, xuất xứ")
            tip.setStyleSheet("color:#92400e; font-weight:600;")
            lay2.addWidget(tip)

            t2 = QTableWidget()
            t2.setColumnCount(8)
            t2.setHorizontalHeaderLabels(["line_id", "STT", "Mã hàng", "Tên hàng", "Model", "Nhãn hiệu", "NSX", "Xuất xứ"])
            t2.setColumnWidth(0, 80)
            t2.setColumnWidth(1, 60)
            t2.setColumnWidth(2, 120)
            t2.setColumnWidth(3, 260)
            t2.setColumnWidth(4, 120)
            t2.setColumnWidth(5, 140)
            t2.setColumnWidth(6, 140)
            t2.setColumnWidth(7, 140)
            t2.setRowCount(len(rows))

            for i, r in enumerate(rows):
                vals = [str(r[0] or ''), str(r[1] or ''), str(r[2] or ''), str(r[3] or ''), str(r[4] or ''), str(r[5] or ''), str(r[6] or ''), str(r[7] or '')]
                for c, v in enumerate(vals):
                    item = QTableWidgetItem(v)
                    if c in (0, 1, 2, 3):
                        item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                    if c in (4, 5, 6, 7) and not v.strip():
                        item.setBackground(Qt.GlobalColor.yellow)
                    t2.setItem(i, c, item)
                t2.setRowHeight(i, 34)

            lay2.addWidget(t2)

            but_save = QPushButton("Lưu bổ sung")

            def _save_verify():
                missing_after = 0
                for i in range(t2.rowCount()):
                    line_id = int(t2.item(i, 0).text())
                    item_code = (t2.item(i, 2).text() or '').strip()
                    model = (t2.item(i, 4).text() or '').strip()
                    brand = (t2.item(i, 5).text() or '').strip()
                    manufacturer = (t2.item(i, 6).text() or '').strip()
                    origin = (t2.item(i, 7).text() or '').strip()

                    if not (model and brand and manufacturer and origin):
                        missing_after += 1

                    misc.sql_commit(
                        "UPDATE fs_supplier_quote_lines SET model=%s, brand=%s, manufacturer=%s, origin=%s WHERE supplier_quote_line_id=%s",
                        (model, brand, manufacturer, origin, line_id),
                    )

                    if item_code:
                        misc.sql_commit(
                            "UPDATE fs_supplier_products SET model=%s, brand=%s, manufacturer=%s, origin=%s, updated_at=NOW() WHERE supplier_id=%s AND item_code=%s",
                            (model, brand, manufacturer, origin, supplier_id, item_code),
                        )

                if missing_after == 0:
                    misc.sql_commit(
                        "UPDATE fs_supplier_files SET process_status='done', processed_by='Anna', processed_at=NOW() WHERE file_id=%s",
                        (file_id,),
                    )
                    QMessageBox.information(vwin, "Xác minh", "Đã bổ sung đủ thông tin. File chuyển sang hoàn tất.")
                else:
                    misc.sql_commit(
                        "UPDATE fs_supplier_files SET process_status='verify_pending', processed_by='Anna', processed_at=NOW() WHERE file_id=%s",
                        (file_id,),
                    )
                    QMessageBox.warning(vwin, "Xác minh", "Vẫn còn dòng thiếu thông tin. File giữ trạng thái Cần xác minh.")

                _load()

            but_save.clicked.connect(_save_verify)
            lay2.addWidget(but_save)

            vwin.setCentralWidget(root2)
            vwin.resize(1100, 560)
            vwin.show()
            self.win_supplier_verify = vwin

        def _load():
            rows = misc.sql_all(
                "SELECT file_id, created_at, file_name, note, process_status, drive_link FROM fs_supplier_files WHERE supplier_id=%s ORDER BY file_id DESC LIMIT 200",
                (supplier_id,),
            ) or []
            tb.setRowCount(len(rows))
            for i, r in enumerate(rows):
                file_id = int(r[0])
                created_at = str(r[1] or '')
                file_name = str(r[2] or '')
                note = str(r[3] or '')
                status = str(r[4] or 'pending')
                drive_link = str(r[5] or '')

                tb.setItem(i, 0, QTableWidgetItem(str(file_id)))
                tb.setItem(i, 1, QTableWidgetItem(created_at))
                tb.setItem(i, 2, QTableWidgetItem(file_name))
                tb.setItem(i, 3, QTableWidgetItem(note))

                status_item = QTableWidgetItem('✅' if status == 'done' else ('🟡 Cần xác minh' if status == 'verify_pending' else ''))
                if status == 'verify_pending':
                    status_item.setBackground(Qt.GlobalColor.yellow)
                tb.setItem(i, 4, status_item)

                if status == 'done':
                    tb.setCellWidget(i, 5, QLabel(''))
                elif status == 'verify_pending':
                    but_verify = QPushButton("Bổ sung")
                    but_verify.clicked.connect(lambda _, fid=file_id: _open_verify(fid))
                    tb.setCellWidget(i, 5, but_verify)
                else:
                    but_import = QPushButton("Nhập vào DB")
                    but_import.clicked.connect(lambda _, fid=file_id, link=drive_link: _import_one(fid, link))
                    tb.setCellWidget(i, 5, but_import)

                but_download = QPushButton("Download")
                but_download.clicked.connect(lambda _, fn=file_name, link=drive_link: _download_one(fn, link))
                tb.setCellWidget(i, 6, but_download)
                tb.setRowHeight(i, 38)

        _load()
        win.setCentralWidget(root)
        win.resize(1120, 620)
        win.show()
        self.win_supplier_quotes = win

    def manage_supplier_products(self, supplier_id: int):
        supplier = misc.sql_one("SELECT supplier_code, name FROM fs_suppliers WHERE supplier_id=%s", (supplier_id,))
        if not supplier:
            QMessageBox.warning(self.win_companyview, "Nhà cung cấp", "Không tìm thấy NCC")
            return

        col_hidden = misc.sql_one(
            "SELECT COUNT(*) FROM information_schema.COLUMNS WHERE TABLE_SCHEMA=DATABASE() AND TABLE_NAME='fs_supplier_products' AND COLUMN_NAME='is_hidden'",
            None,
        )
        if not col_hidden or int(col_hidden[0] or 0) == 0:
            misc.sql_commit("ALTER TABLE fs_supplier_products ADD COLUMN is_hidden TINYINT(1) NOT NULL DEFAULT 0", None)

        rows = misc.sql_all(
            "SELECT item_code, item_name, model, brand, manufacturer, origin, latest_price FROM fs_supplier_products WHERE supplier_id=%s AND IFNULL(is_hidden,0)=0 ORDER BY id DESC",
            (supplier_id,),
        ) or []

        lines = [f"NCC {supplier[0]} - {supplier[1]}", "", "Danh sách mã hàng hiện có:"]
        if rows:
            for x in rows:
                lines.append(
                    f"- {x[0]} | {x[1] or ''} | model: {x[2] or ''} | nhãn hiệu: {x[3] or ''} | NSX: {x[4] or ''} | xuất xứ: {x[5] or ''} | giá {int(x[6] or 0):,}"
                )
        else:
            lines.append("(chưa có mã hàng)")

        msg = QMessageBox(self.win_companyview)
        msg.setWindowTitle("Mã hàng NCC")
        msg.setText("\n".join(lines))
        but_add = msg.addButton("Thêm/Cập nhật", QMessageBox.ButtonRole.AcceptRole)
        msg.addButton("Đóng", QMessageBox.ButtonRole.RejectRole)
        msg.exec()

        if msg.clickedButton() != but_add:
            return

        item_code, ok = QInputDialog.getText(self.win_companyview, "Mã hàng NCC", "Mã hàng (item_code):")
        if not ok or not str(item_code).strip():
            return
        item_name, _ = QInputDialog.getText(self.win_companyview, "Mã hàng NCC", "Tên hàng:")
        model, _ = QInputDialog.getText(self.win_companyview, "Mã hàng NCC", "Model:", text=item_code.strip())
        brand, _ = QInputDialog.getText(self.win_companyview, "Mã hàng NCC", "Nhãn hiệu:")
        manufacturer, _ = QInputDialog.getText(self.win_companyview, "Mã hàng NCC", "Nhà sản xuất:")
        origin, _ = QInputDialog.getText(self.win_companyview, "Mã hàng NCC", "Xuất xứ:")
        latest_price, _ = QInputDialog.getDouble(self.win_companyview, "Mã hàng NCC", "Giá nhập gần nhất:", 0, 0, 10**12, 0)

        exists = misc.sql_one(
            "SELECT id FROM fs_supplier_products WHERE supplier_id=%s AND item_code=%s",
            (supplier_id, item_code.strip()),
        )
        if exists:
            misc.sql_commit(
                "UPDATE fs_supplier_products SET item_name=%s, model=%s, brand=%s, manufacturer=%s, origin=%s, latest_price=%s, is_hidden=0 WHERE supplier_id=%s AND item_code=%s",
                (item_name.strip(), model.strip(), brand.strip(), manufacturer.strip(), origin.strip(), latest_price, supplier_id, item_code.strip()),
            )
        else:
            misc.sql_commit(
                "INSERT INTO fs_supplier_products (supplier_id, item_code, item_name, model, brand, manufacturer, origin, latest_price) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
                (supplier_id, item_code.strip(), item_name.strip(), model.strip(), brand.strip(), manufacturer.strip(), origin.strip(), latest_price),
            )

        QMessageBox.information(self.win_companyview, "Mã hàng NCC", "Đã lưu thông tin mã hàng NCC")

    def on_tab_changed(self, index):
        self.uic8.tableWidget_2.clear()

        if self.suppliers_tab is not None and index == self.uic8.tabWidget.indexOf(self.suppliers_tab):
            self.uic8.label.setText('DANH SÁCH NHÀ CUNG CẤP')
            self.load_suppliers((self.suppliers_search.text() or '').strip() if self.suppliers_search else '')
            return

        if index == self.uic8.tabWidget.indexOf(self.uic8.tab_2):
            # self.uic8.tableWidget_2.verticalHeader().setVisible(False)
            self.uic8.label.setText('DANH SÁCH ĐỐI TÁC CÁ NHÂN')

            # Tối ưu hiệu năng: chỉ load dữ liệu gần nhất khi mở tab Cá nhân
            result = misc.sql_all("SELECT * FROM ds_ca_nhan ORDER BY dien_thoai DESC LIMIT 500", None)

            if result:
                self.uic8.tableWidget_2.setColumnCount(5)
                self.uic8.tableWidget_2.setRowCount(len(result))
                self.uic8.tableWidget_2.setHorizontalHeaderLabels(
                    ['Tên khách hàng', 'Điện thoại', 'Số lead', 'Công ty', 'Thao tác'])
                self.uic8.tableWidget_2.setColumnWidth(0, 130)
                self.uic8.tableWidget_2.setColumnWidth(1, 90)
                self.uic8.tableWidget_2.setColumnWidth(2, 60)
                self.uic8.tableWidget_2.setColumnWidth(3, 210)
                self.uic8.tableWidget_2.setColumnWidth(4, 70)

                for row in range(len(result)):
                    self.uic8.tableWidget_2.setItem(row, 0, QTableWidgetItem(str(result[row][0])))
                    self.uic8.tableWidget_2.setItem(row, 1, QTableWidgetItem(str(result[row][1])))
                    if result[row][2] is not None:
                        so_lead = str(len(result[row][2].split('|')))
                    else:
                        so_lead = '0'
                    item = QTableWidgetItem(so_lead)
                    item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    self.uic8.tableWidget_2.setItem(row, 2, item)

                    if result[row][6] is not None:
                        ten_cty = result[row][6]
                    else:
                        ten_cty = ' '

                    if result[row][7] is not None:
                        mst = 'MST: ' + result[row][7]
                    else:
                        mst = ' '
                    txt = ten_cty + '\n' + mst

                    self.uic8.tableWidget_2.setItem(row, 3, QTableWidgetItem(txt))

                    self.uic8.tableWidget_2.resizeRowToContents(row)

                    but1 = QPushButton('Xem')
                    but1.clicked.connect(lambda _, r=row: Crm.view_detail_canhan(self, r))

                    self.uic8.tableWidget_2.setCellWidget(row, 4, but1)

                    self.uic8.tableWidget_2.setRowHeight(row, 70)

                self.uic8.tableWidget_2.repaint()
            else:
                print('Lỗi không hiển thị được')
                self.uic8.tableWidget_2.clear()
                pass

            # Không connect lại signal ở đây để tránh click 1 lần chạy nhiều lần
            if self.uic8.textEdit.toPlainText() != '':
                Crm.search_company(self)

        else:
            self.uic8.label.setText('DANH SÁCH CÔNG TY ĐỐI TÁC')
            self.uic8.tableWidget.verticalHeader().setVisible(False)

            code = "SELECT * FROM ds_cong_ty LIMIT 100"
            result = misc.sql_all(code, None)

            if result:
                self.uic8.tableWidget.setColumnCount(5)
                self.uic8.tableWidget.setRowCount(len(result))
                self.uic8.tableWidget.setHorizontalHeaderLabels(
                    ['Mã KH', 'Mã số thuế', 'Tên công ty', 'Địa chỉ', 'Thao tác'])
                self.uic8.tableWidget.setColumnWidth(0, 40)
                self.uic8.tableWidget.setColumnWidth(1, 80)
                self.uic8.tableWidget.setColumnWidth(2, 200)
                self.uic8.tableWidget.setColumnWidth(3, 220)
                self.uic8.tableWidget.setColumnWidth(4, 50)

                for row in range(len(result)):
                    self.uic8.tableWidget.setItem(row, 0, QTableWidgetItem(str(result[row][0])))
                    self.uic8.tableWidget.setItem(row, 1, QTableWidgetItem(result[row][2]))
                    self.uic8.tableWidget.setItem(row, 2, QTableWidgetItem(result[row][1]))
                    self.uic8.tableWidget.setItem(row, 3, QTableWidgetItem(result[row][4]))

                    but1 = QPushButton('Xem')
                    but1.clicked.connect(lambda _, r=row: Crm.view_detail_company(self, result[r][2]))
                    self.uic8.tableWidget.setCellWidget(row, 4, but1)

                    self.uic8.tableWidget.setRowHeight(row, 70)

                self.uic8.tableWidget.repaint()
            else:
                print('Lỗi không hiển thị được')
                self.uic8.tableWidget.clear()
                pass

            if self.uic8.textEdit.toPlainText() != '':
                Crm.search_company(self)

    def _open_new_lead_screen(self):
        """Mở màn hình tạo lead mới từ màn CRM chi tiết công ty/cá nhân."""
        try:
            if hasattr(self, 'lead_handler') and self.lead_handler:
                # Đồng bộ user nếu có trên main window
                if hasattr(self, 'user'):
                    self.lead_handler.user = self.user
                if hasattr(self, 'user_phone'):
                    self.lead_handler.user_phone = self.user_phone
                self.lead_handler.create_new_lead()
                return

            # Fallback: khởi tạo tạm LeadHandle nếu chưa có trên context hiện tại
            from lead_handle import LeadHandle
            self.lead_handler = LeadHandle(self)
            if hasattr(self, 'user'):
                self.lead_handler.user = self.user
            if hasattr(self, 'user_phone'):
                self.lead_handler.user_phone = self.user_phone
            self.lead_handler.create_new_lead()
        except Exception as e:
            QMessageBox.warning(None, "Lỗi", f"Không mở được màn hình tạo lead mới: {e}")

    def _open_quote_review_for_contract(self):
        """Mở danh sách báo giá cũ để user chọn xem lại theo luồng tạo hợp đồng."""
        try:
            mst = getattr(self, '_current_mst', '') or ''
            phone = getattr(self, '_current_phone', '') or ''

            data = Crm._history_data(self, mst=mst, phone=phone)
            quotes = data.get('quotes', []) or []

            if not quotes:
                QMessageBox.information(None, "Báo giá", "Không tìm thấy báo giá cũ cho khách hàng này.")
                return

            self.win_quote_picker = QMainWindow()
            self.win_quote_picker.setWindowTitle("Chọn báo giá để xem lại")
            container = QWidget(self.win_quote_picker)
            layout = QVBoxLayout(container)
            table = QTableWidget(container)
            table.setColumnCount(6)
            table.setRowCount(len(quotes))
            table.setHorizontalHeaderLabels(['Số BG', 'Ngày', 'Tiêu đề', 'Giá trị', 'Sale', 'Thao tác'])
            table.setColumnWidth(0, 70)
            table.setColumnWidth(1, 90)
            table.setColumnWidth(2, 260)
            table.setColumnWidth(3, 120)
            table.setColumnWidth(4, 110)
            table.setColumnWidth(5, 100)
            table.verticalHeader().setVisible(False)

            for i, q in enumerate(quotes):
                so_bg = int(q[0])
                ngay = str(q[2] or '')
                tieu_de = str(q[6] or '')
                gia_tri = int(q[11] or 0) if len(q) > 11 and q[11] is not None else 0
                user_bg = str(q[5] or '') if len(q) > 5 else ''

                table.setItem(i, 0, QTableWidgetItem(str(so_bg)))
                table.setItem(i, 1, QTableWidgetItem(ngay))
                table.setItem(i, 2, QTableWidgetItem(tieu_de[:120]))
                table.setItem(i, 3, QTableWidgetItem("{:,.0f}".format(gia_tri)))
                table.setItem(i, 4, QTableWidgetItem(user_bg))

                btn = QPushButton('Xem lại')
                btn.clicked.connect(lambda _, sb=so_bg: Crm._open_quote_by_so_bg(self, sb))
                table.setCellWidget(i, 5, btn)
                table.setRowHeight(i, 58)

            layout.addWidget(table)
            self.win_quote_picker.setCentralWidget(container)
            self.win_quote_picker.resize(860, 520)
            self.win_quote_picker.show()

        except Exception as e:
            QMessageBox.warning(None, "Lỗi", f"Không mở được danh sách báo giá: {e}")

    def _open_quote_by_so_bg(self, so_bg):
        """Mở màn hình xem lại báo giá theo số BG, từ đó đi tiếp luồng tạo hợp đồng."""
        try:
            row = misc.sql_one("SELECT * FROM ds_bao_gia WHERE so_bg = %s", (so_bg,))
            if not row:
                QMessageBox.information(None, "Báo giá", f"Không tìm thấy báo giá #{so_bg}")
                return

            lead_id = row[1]
            if not lead_id:
                QMessageBox.warning(None, "Lỗi", f"Báo giá #{so_bg} chưa gắn lead_id.")
                return

            data = None
            if row[3]:
                goods = str(row[3]).split('@')
                data = [item.split('|') for item in goods if item]

            self.win_quotato = quotation.Quotato()
            if hasattr(self, 'user'):
                self.win_quotato.user = self.user
            if hasattr(self, 'user_phone'):
                self.win_quotato.user_phone = self.user_phone

            # Tạo cửa sổ báo giá và đổ nội dung trực tiếp
            self.win_quotato.sub_win1 = QMainWindow()
            self.win_quotato.uic5 = quotation.Ui_Win_bao_gia()
            self.win_quotato.uic5.setupUi(self.win_quotato.sub_win1)
            apply_ui_v2(self.win_quotato.sub_win1)
            if hasattr(quotation.Quotato, '_polish_grid_buttons'):
                quotation.Quotato._polish_grid_buttons(self.win_quotato.uic5)
            self.win_quotato.sub_win1.show()
            self.win_quotato.show_bg(str(lead_id), str(so_bg), data)

        except Exception as e:
            QMessageBox.warning(None, "Lỗi", f"Không mở được báo giá #{so_bg}: {e}")

    def view_detail_canhan(self, row_index):
        item = self.uic8.tableWidget_2.item(row_index, 1)

        # Hiển thị form chi tiết
        self.win_detailcompanyview = QMainWindow()
        self.uic10 = Ui_ViewDetailCompany()
        self.uic10.setupUi(self.win_detailcompanyview)
        apply_ui_v2(self.win_detailcompanyview)
        self.win_detailcompanyview.show()

        if not item:
            self.uic10.label_noti.setText("⚠️ Không có dữ liệu số điện thoại tại dòng này.")
            return

        dien_thoai = item.text().strip()
        self._current_phone = dien_thoai
        self._current_mst = ''
        if not dien_thoai:
            self.uic10.label_noti.setText("⚠️ Số điện thoại rỗng.")
            return

        kq = misc.sql_one("SELECT * FROM ds_ca_nhan WHERE dien_thoai = %s", (dien_thoai,))
        if kq and len(kq) > 7 and kq[7]:
            self._current_mst = str(kq[7]).strip()
        if not kq:
            self.uic10.label_noti.setText("❌ Không tìm thấy dữ liệu cá nhân trong CSDL.")
            return

        self.uic10.but_update.clicked.connect(lambda: Crm.update_thong_tin_ca_nhan(self))
        self.uic10.but_tao_lead.clicked.connect(lambda: Crm._open_new_lead_screen(self))
        self.uic10.but_tao_hop_dong.clicked.connect(lambda: Crm._open_quote_review_for_contract(self))
        self.uic10.label_noti.setText(kq[0])
        self.uic10.but_tao_moi.hide()

        # Lịch sử giao dịch: load khi người dùng mở tab
        try:
            self.uic10.tabWidget.currentChanged.disconnect()
        except Exception:
            pass
        mst_canhan = kq[7] if len(kq) > 7 and kq[7] else ''
        self.uic10.tabWidget.currentChanged.connect(
            lambda idx, _mst=mst_canhan, _phone=dien_thoai: Crm._render_history_tab(self, self.uic10, mst=_mst, phone=_phone)
            if idx == self.uic10.tabWidget.indexOf(self.uic10.tab_3) else None
        )

        self.uic10.txt_nguoi_lien_he.setText(kq[0])
        self.uic10.txt_sdt_nguoi_lien_he.setText(kq[1])

        if kq[5]:
            self.uic10.txt_email_nguoi_lien_he.setText(kq[5])
        else:
            self.uic10.txt_email_nguoi_lien_he.setPlaceholderText('Email cá nhân đang trống')

        if kq[7]:
            cong_ty = misc.sql_one("SELECT * FROM ds_cong_ty WHERE mst = %s", (kq[7],))
            if cong_ty:
                self.uic10.txt_ten_cong_ty.setText(cong_ty[1])
                self.uic10.txt_dia_chi.setText(cong_ty[4] if cong_ty[4] else '')
                self.uic10.txt_mst.setText(cong_ty[2])
                self.uic10.txt_sdt_cong_ty.setText(cong_ty[7])

        if kq[6]:
            self.uic10.txt_nguoi_phu_trach.setText(kq[6])

    def show_canhan_lichsugiaodich(self):
        dien_thoai = self.uic9.txt_sdt_nguoi_lien_he.toPlainText()

        kq = misc.sql_all("SELECT * FROM ds_bao_gia WHERE dien_thoai = %s", (dien_thoai,))

        if kq:
            self.uic9.tableWidget.setColumnCount(6)
            self.uic9.tableWidget.setRowCount(len(kq))
            self.uic9.tableWidget.setHorizontalHeaderLabels(
                ['Số BG', 'Ngày tháng', 'Tổng tiền', 'Người làm BG', 'Thành công'])
            self.uic9.tableWidget.setColumnWidth(0, 50)
            self.uic9.tableWidget.setColumnWidth(1, 80)
            self.uic9.tableWidget.setColumnWidth(2, 100)
            self.uic9.tableWidget.setColumnWidth(3, 120)
            self.uic9.tableWidget.setColumnWidth(4, 120)
            self.uic9.tableWidget.setColumnWidth(5, 120)

            # Function to convert date string to QDate
            def date_str_to_qdate(date_str):
                return QDate.fromString(date_str, "dd/MM/yy")

            # Sort the list based on the converted QDate, newer to older
            kq.sort(key=lambda x: date_str_to_qdate(x[2]), reverse=True)

            for row in range(len(kq)):
                self.uic9.tableWidget.setItem(row, 0, QTableWidgetItem(str(kq[row][0])))
                self.uic9.tableWidget.setItem(row, 1, QTableWidgetItem(str(kq[row][2])))

                if kq[row][11] is not None:
                    temp = "{:,.0f}".format((kq[row][11]))
                    self.uic9.tableWidget.setItem(row, 2, QTableWidgetItem(temp))

                if kq[row][5] is not None:
                    self.uic9.tableWidget.setItem(row, 3, QTableWidgetItem(kq[row][5]))

                if kq[row][4] is not None:
                    if kq[row][4] == "N":
                        self.uic9.tableWidget.setItem(row, 4, QTableWidgetItem("Chưa thành công"))
                    elif kq[row][4] == "T" or kq[row][4] == '1':
                        self.uic9.tableWidget.setItem(row, 4, QTableWidgetItem("Thành công"))

                but1 = QPushButton('Xem lại')
                # but1.clicked.connect(lambda: file_upload.download_file(self.uic7.tableWidget.item(self.uic7.tableWidget.currentRow(), 0).text()))
                self.uic9.tableWidget.setCellWidget(row, 5, but1)

                self.uic9.tableWidget.setRowHeight(row, 65)

                self.uic9.txt_nguoi_phu_trach.setText(kq[0][5])

        else:
            print('Lỗi không hiển thị được')
            self.uic9.tableWidget.clear()
            pass

    def _load_company_contacts(self, mst):
        mst = (mst or '').strip()
        if not mst:
            return []
        rows = misc.sql_all(
            "SELECT ten, dien_thoai, email FROM ds_ca_nhan WHERE mst_cong_ty = %s ORDER BY dien_thoai DESC",
            (mst,)
        ) or []
        out = []
        seen = set()
        for r in rows:
            ten = str(r[0] or '').strip()
            sdt = str(r[1] or '').strip()
            email = str(r[2] or '').strip()
            key = (ten, sdt, email)
            if key in seen:
                continue
            seen.add(key)
            out.append((ten, sdt, email))
        return out

    def _render_company_contacts_note(self, mst, company_name=''):
        contacts = Crm._load_company_contacts(self, mst)
        if not contacts:
            self.uic9.label_noti.setText(company_name or '')
            return

        lines = [f"{company_name}".strip(), f"Liên hệ công ty: {len(contacts)} người"]
        for i, (ten, sdt, email) in enumerate(contacts, start=1):
            item = f"{i}) {ten or '(chưa tên)'} - {sdt or '(chưa SĐT)'}"
            if email:
                item += f" - {email}"
            lines.append(item)
        self.uic9.label_noti.setText("\n".join([x for x in lines if x]))

    def view_detail_company(self, mst, so_bg=None):
        # Khởi tạo màn hình
        self.win_detailcompanyview = QMainWindow()
        self.uic9 = Ui_ViewDetailCompany()
        self.uic9.setupUi(self.win_detailcompanyview)
        apply_ui_v2(self.win_detailcompanyview)
        self.win_detailcompanyview.show()

        # Phân ra 2 trường hợp để lấy MST
        if isinstance(mst, str):
            mst = mst.strip()
        else:
            row = self.uic8.tableWidget.currentRow()
            item_mst = self.uic8.tableWidget.item(row, 1) if row >= 0 else None
            mst = item_mst.text().strip() if item_mst else ''

        if not mst:
            self.uic9.label_noti.setText("❌ Không xác định được MST công ty.")
            return

        self._current_mst = mst

        kq = misc.sql_one("SELECT * FROM ds_cong_ty WHERE mst = %s", (mst,))
        if not kq:
            self.uic9.label_noti.setText("❌ Không tìm thấy công ty theo MST.")
            return

        Crm._render_company_contacts_note(self, mst, kq[1])
        if so_bg:
            self.uic9.label_noti.setText(self.uic9.label_noti.text() + '\n' + 'Soạn hợp đồng theo báo giá số ' + str(so_bg))
            self.uic9.label_so_hd.setText(so_bg)

        self.uic9.label_ma_kh.setText(str(kq[0]))
        self.uic9.but_tao_moi.hide()

        self.uic9.txt_ten_cong_ty.setText(kq[1])
        self.uic9.txt_mst.setText(kq[2])

        if kq[4] is None:
            self.uic9.txt_dia_chi.setPlaceholderText('Chưa ghi địa chỉ công ty')
        else:
            self.uic9.txt_dia_chi.setText(kq[4])

        self.uic9.txt_nguoi_lien_he.setText(kq[6])

        self._current_phone = (kq[7] or '').strip() if len(kq) > 7 and kq[7] is not None else ''

        if kq[7] is None:
            self.uic9.txt_sdt_cong_ty.setPlaceholderText('Chưa ghi số điện thoại công ty')
        else:
            self.uic9.txt_sdt_cong_ty.setText(kq[7])

        if kq[8] is None:
            self.uic9.txt_email.setPlaceholderText('Chưa ghi email công ty')
        else:
            self.uic9.txt_email.setText(kq[8])

        self.uic9.txt_nguoi_phu_trach.setText(kq[9])

        if kq[10] is None:
            self.uic9.txt_nguoi_dai_dien.setPlaceholderText('Chưa có tên người đại diện công ty')
        else:
            self.uic9.txt_nguoi_dai_dien.setText(kq[10])

        if kq[11] is None:
            self.uic9.txt_stk.setPlaceholderText('Chưa có số tài khoản công ty')
        else:
            self.uic9.txt_stk.setText(kq[11])

        if kq[12] is None:
            self.uic9.txt_stk.setPlaceholderText('Tại ngân hàng nào???')
        else:
            self.uic9.txt_ten_ngan_hang.setText(kq[12])

        if kq[13] is None:
            self.uic9.txt_chuc_vu.setPlaceholderText('Cần ghi rõ chức vụ của người đại diện')
        else:
            self.uic9.txt_chuc_vu.setText(kq[13])

        if kq[14] is None:
            self.uic9.txt_sdt_nguoi_dai_dien.setPlaceholderText('Số điện thoại của người đại diện công ty')
        else:
            self.uic9.txt_sdt_nguoi_dai_dien.setText(kq[14])

        if kq[15] is None:
            self.uic9.txt_chuc_vu_nguoi_lien_he.setPlaceholderText('Chức vụ của người liên hệ')
        else:
            self.uic9.txt_chuc_vu_nguoi_lien_he.setText(kq[15])

        if kq[16] is None:
            self.uic9.txt_email_nguoi_dai_dien.setPlaceholderText('Email của người đại diện công ty')
        else:
            self.uic9.txt_email_nguoi_dai_dien.setText(kq[16])

        self.uic9.txt_sdt_nguoi_lien_he.setText(kq[17])

        if kq[18] is None:
            self.uic9.txt_email_nguoi_lien_he.setPlaceholderText('Email của người liên hệ')
        else:
            self.uic9.txt_email_nguoi_lien_he.setText(kq[18])

        self.uic9.but_update.clicked.connect(lambda: Crm.update_tt_cong_ty(self))
        self.uic9.but_tao_lead.clicked.connect(lambda: Crm._open_new_lead_screen(self))
        if so_bg:
            self.uic9.but_tao_hop_dong.clicked.connect(lambda: Crm.tao_hop_dong_tu_mau(self))
        else:
            self.uic9.but_tao_hop_dong.clicked.connect(lambda: Crm._open_quote_review_for_contract(self))

        # Lịch sử giao dịch: load khi mở tab "Lịch sử giao dịch"
        try:
            self.uic9.tabWidget.currentChanged.disconnect()
        except Exception:
            pass
        phone_company = kq[7] if len(kq) > 7 and kq[7] else ''
        self.uic9.tabWidget.currentChanged.connect(
            lambda idx, _mst=mst, _phone=phone_company: Crm._render_history_tab(self, self.uic9, mst=_mst, phone=_phone)
            if idx == self.uic9.tabWidget.indexOf(self.uic9.tab_3) else None
        )

    def update_tt_cong_ty(self):
        # B1: Thu thập dữ liệu từ giao diện
        fields = {
            "ten_cty": self.uic9.txt_ten_cong_ty.toPlainText().strip(),
            "dia_chi": self.uic9.txt_dia_chi.toPlainText().strip(),
            "mst": self.uic9.txt_mst.toPlainText().strip(),
            "email_cty": self.uic9.txt_email.toPlainText().strip(),
            "stk": self.uic9.txt_stk.toPlainText().strip(),
            "ngan_hang": self.uic9.txt_ten_ngan_hang.toPlainText().strip(),
            "sdt_cty": self.uic9.txt_sdt_cong_ty.toPlainText().strip(),
            "nguoi_dd": self.uic9.txt_nguoi_dai_dien.toPlainText().strip(),
            "sdt_ndd": self.uic9.txt_sdt_nguoi_dai_dien.toPlainText().strip(),
            "chuc_vu": self.uic9.txt_chuc_vu.toPlainText().strip(),
            "email_ndd": self.uic9.txt_email_nguoi_dai_dien.toPlainText().strip(),
            "nguoi_lh": self.uic9.txt_nguoi_lien_he.toPlainText().strip(),
            "sdt_nlh": self.uic9.txt_sdt_nguoi_lien_he.toPlainText().strip(),
            "chuc_vu_nlh": self.uic9.txt_chuc_vu_nguoi_lien_he.toPlainText().strip(),
            "email_nlh": self.uic9.txt_email_nguoi_lien_he.toPlainText().strip(),
        }

        # B1.1: Chuẩn hoá dữ liệu đầu vào (giảm bẩn/trùng)
        fields["mst"] = Crm.normalize_mst(fields["mst"])
        fields["sdt_cty"] = Crm.normalize_phone(fields["sdt_cty"])
        fields["sdt_ndd"] = Crm.normalize_phone(fields["sdt_ndd"])
        fields["sdt_nlh"] = Crm.normalize_phone(fields["sdt_nlh"])

        # Đồng bộ lại UI sau chuẩn hoá để user nhìn thấy dữ liệu sạch
        self.uic9.txt_mst.setText(fields["mst"])
        self.uic9.txt_sdt_cong_ty.setText(fields["sdt_cty"])
        self.uic9.txt_sdt_nguoi_dai_dien.setText(fields["sdt_ndd"])
        self.uic9.txt_sdt_nguoi_lien_he.setText(fields["sdt_nlh"])

        # B2: Kiểm tra bắt buộc
        required_fields = {
            "Tên công ty": fields["ten_cty"],
            "Mã số thuế": fields["mst"],
            "Địa chỉ": fields["dia_chi"],
            "Người đại diện": fields["nguoi_dd"],
            "Điện thoại công ty": fields["sdt_cty"],
        }

        missing_required = [label for label, value in required_fields.items() if not value]

        # Nếu thiếu trường bắt buộc nhưng có người liên hệ → vẫn lưu người liên hệ vào ds_ca_nhan
        if missing_required:
            required_snapshot = {
                "Tên công ty": fields["ten_cty"],
                "Mã số thuế": fields["mst"],
                "Địa chỉ": fields["dia_chi"],
                "Người đại diện": fields["nguoi_dd"],
                "Điện thoại công ty": fields["sdt_cty"],
            }
            print("[update_tt_cong_ty] required field snapshot:", required_snapshot)

            if fields["nguoi_lh"] and re.match(r"^0\d{9}$", fields["sdt_nlh"]):
                misc.sql_commit(
                    """
                    INSERT INTO ds_ca_nhan (ten, dien_thoai, email, ten_cong_ty, mst_cong_ty)
                    VALUES (%s, %s, %s, %s, %s)
                    ON DUPLICATE KEY UPDATE
                        ten = VALUES(ten),
                        email = VALUES(email),
                        ten_cong_ty = VALUES(ten_cong_ty),
                        mst_cong_ty = VALUES(mst_cong_ty)
                    """,
                    (fields["nguoi_lh"], fields["sdt_nlh"], fields["email_nlh"], fields["ten_cty"], fields["mst"])
                )
                self.uic9.label_noti.setStyleSheet("color: orange")
                self.uic9.label_noti.setText(
                    "⚠️ Thiếu thông tin công ty: " + ", ".join(missing_required) + ". Chỉ lưu người liên hệ."
                )
                return
            else:
                self.uic9.label_noti.setStyleSheet("color: red")
                self.uic9.label_noti.setText(f"❌ Thiếu thông tin: {', '.join(missing_required)}")
                return

        # B3: Kiểm tra định dạng mã số thuế
        mst = fields["mst"]
        if not re.match(r"^\d{10}(-\d{3})?$", mst):
            self.uic9.label_noti.setStyleSheet("color: red")
            self.uic9.label_noti.setText("❌ Mã số thuế không đúng định dạng! (VD: 1234567890 hoặc 1234567890-001)")
            return

        # B3.1: Validate định dạng số điện thoại sau chuẩn hoá
        if not re.match(r"^0\d{9}$", fields["sdt_cty"]):
            self.uic9.label_noti.setStyleSheet("color: red")
            self.uic9.label_noti.setText("❌ Số điện thoại công ty phải hợp lệ (10 số, bắt đầu bằng 0).")
            return

        for label, phone_key in [("SĐT người đại diện", "sdt_ndd"), ("SĐT người liên hệ", "sdt_nlh")]:
            if fields[phone_key] and not re.match(r"^0\d{9}$", fields[phone_key]):
                self.uic9.label_noti.setStyleSheet("color: red")
                self.uic9.label_noti.setText(f"❌ {label} không hợp lệ.")
                return

        # B4: Gán biến tắt
        ten_cty = fields["ten_cty"]
        dia_chi = fields["dia_chi"]
        email_cty = fields["email_cty"]
        stk = fields["stk"]
        ngan_hang = fields["ngan_hang"]
        sdt_cty = fields["sdt_cty"]
        nguoi_dd = fields["nguoi_dd"]
        sdt_ndd = fields["sdt_ndd"]
        chuc_vu = fields["chuc_vu"]
        email_ndd = fields["email_ndd"]
        nguoi_lh = fields["nguoi_lh"]
        sdt_nlh = fields["sdt_nlh"]
        chuc_vu_nlh = fields["chuc_vu_nlh"]
        email_nlh = fields["email_nlh"]
        ma_kh = self.uic9.label_ma_kh.text()

        try:
            # B5: Cập nhật hoặc thêm công ty
            kq = misc.sql_one("SELECT * from ds_cong_ty WHERE mst = %s", (mst,))
            if kq:
                code = """UPDATE ds_cong_ty SET ten_cong_ty = %s, dia_chi = %s, mst = %s, email_cong_ty = %s, stk = %s, 
                          ngan_hang = %s, dien_thoai_cong_ty = %s, nguoi_dai_dien = %s, sdt_nguoi_dd = %s, 
                          chuc_vu_ndd = %s, email_ndd = %s, nguoi_lien_he = %s, sdt_nguoi_lh = %s, chuc_vu_nlh = %s, 
                          email_nlh = %s WHERE mst = %s"""
                params = (
                    ten_cty, dia_chi, mst, email_cty, stk, ngan_hang, sdt_cty,
                    nguoi_dd, sdt_ndd, chuc_vu, email_ndd,
                    nguoi_lh, sdt_nlh, chuc_vu_nlh, email_nlh, mst
                )
                misc.sql_commit(code, params)
                self.uic9.label_noti.setStyleSheet("color: green")
                self.uic9.label_noti.setText("✅ Đã cập nhật thông tin công ty.")
            else:
                code = """INSERT INTO ds_cong_ty 
                          (ten_cong_ty, mst, dia_chi, email_cong_ty, stk, ngan_hang, dien_thoai_cong_ty, 
                           nguoi_dai_dien, sdt_nguoi_dd, chuc_vu_ndd, email_ndd, nguoi_lien_he, 
                           sdt_nguoi_lh, chuc_vu_nlh, email_nlh)
                          VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"""
                params = (
                    ten_cty, mst, dia_chi, email_cty, stk, ngan_hang, sdt_cty,
                    nguoi_dd, sdt_ndd, chuc_vu, email_ndd,
                    nguoi_lh, sdt_nlh, chuc_vu_nlh, email_nlh
                )
                misc.sql_commit(code, params)
                self.uic9.label_noti.setStyleSheet("color: green")
                self.uic9.label_noti.setText("✅ Đã thêm công ty mới vào cơ sở dữ liệu.")

            # B6: Upsert cá nhân (người đại diện) - atomic, tránh race duplicate
            if nguoi_dd.strip() and re.match(r"^0\d{9}$", sdt_ndd):
                misc.sql_commit(
                    """
                    INSERT INTO ds_ca_nhan (ten, dien_thoai, email, ten_cong_ty, mst_cong_ty)
                    VALUES (%s, %s, %s, %s, %s)
                    ON DUPLICATE KEY UPDATE
                        ten = VALUES(ten),
                        email = VALUES(email),
                        ten_cong_ty = VALUES(ten_cong_ty),
                        mst_cong_ty = VALUES(mst_cong_ty)
                    """,
                    (nguoi_dd, sdt_ndd, email_ndd, ten_cty, mst)
                )

            # B7: Upsert cá nhân (người liên hệ) - atomic, tránh race duplicate
            if nguoi_lh.strip() and re.match(r"^0\d{9}$", sdt_nlh):
                misc.sql_commit(
                    """
                    INSERT INTO ds_ca_nhan (ten, dien_thoai, email, ten_cong_ty, mst_cong_ty)
                    VALUES (%s, %s, %s, %s, %s)
                    ON DUPLICATE KEY UPDATE
                        ten = VALUES(ten),
                        email = VALUES(email),
                        ten_cong_ty = VALUES(ten_cong_ty),
                        mst_cong_ty = VALUES(mst_cong_ty)
                    """,
                    (nguoi_lh, sdt_nlh, email_nlh, ten_cty, mst)
                )

            # B8: Hiển thị đầy đủ danh sách người liên hệ của công ty trong CRM
            Crm._render_company_contacts_note(self, mst, ten_cty)

        except Exception as e:
            self.uic9.label_noti.setStyleSheet("color: red")
            self.uic9.label_noti.setText(f"❌ Lỗi khi cập nhật thông tin: {str(e)}")
            return

    def update_thong_tin_ca_nhan(self):
        # B1: Lấy dữ liệu từ giao diện
        ten = self.uic10.txt_nguoi_lien_he.toPlainText().strip()
        dien_thoai = Crm.normalize_phone(self.uic10.txt_sdt_nguoi_lien_he.toPlainText().strip())
        email = self.uic10.txt_email_nguoi_lien_he.toPlainText().strip()
        ten_cong_ty = self.uic10.txt_ten_cong_ty.toPlainText().strip()
        mst = Crm.normalize_mst(self.uic10.txt_mst.toPlainText().strip())

        self.uic10.txt_sdt_nguoi_lien_he.setText(dien_thoai)
        self.uic10.txt_mst.setText(mst)

        # B2: Kiểm tra bắt buộc
        if not ten or not dien_thoai:
            self.uic10.label_noti.setStyleSheet("color: red")
            self.uic10.label_noti.setText("❌ Vui lòng nhập đầy đủ tên và số điện thoại")
            return

        # B3: Kiểm tra định dạng số điện thoại (10 số VN, bắt đầu bằng 0)
        if not re.match(r"^0\d{9}$", dien_thoai):
            self.uic10.label_noti.setStyleSheet("color: red")
            self.uic10.label_noti.setText("❌ Số điện thoại không hợp lệ (10 số, bắt đầu bằng 0).")
            return

        if mst and not re.match(r"^\d{10}(-\d{3})?$", mst):
            self.uic10.label_noti.setStyleSheet("color: red")
            self.uic10.label_noti.setText("❌ Mã số thuế không đúng định dạng!")
            return

        # B4: Upsert vào bảng ds_ca_nhan (trùng SĐT thì cập nhật)
        misc.sql_commit(
            """
            INSERT INTO ds_ca_nhan (ten, dien_thoai, email, ten_cong_ty, mst_cong_ty)
            VALUES (%s, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                ten = VALUES(ten),
                email = VALUES(email),
                ten_cong_ty = VALUES(ten_cong_ty),
                mst_cong_ty = VALUES(mst_cong_ty)
            """,
            (ten, dien_thoai, email, ten_cong_ty, mst)
        )

        # B5: Hiển thị thông báo
        self.uic10.label_noti.setStyleSheet("color: green")
        self.uic10.label_noti.setText("✅ Đã cập nhật thông tin cá nhân.")

    def tao_hop_dong_tu_mau(self):
        try:
            so_bg = (self.uic9.label_so_hd.text() or '').strip()
            ten_cty = self.uic9.txt_ten_cong_ty.toPlainText().strip()
            mst = self.uic9.txt_mst.toPlainText().strip()
            dia_chi = self.uic9.txt_dia_chi.toPlainText().strip()
            sdt = self.uic9.txt_sdt_nguoi_dai_dien.toPlainText().strip() or self.uic9.txt_sdt_cong_ty.toPlainText().strip()

            if not so_bg:
                self.uic9.label_noti.setStyleSheet('color: red')
                self.uic9.label_noti.setText('❌ Chưa có số báo giá để tạo hợp đồng.')
                return

            base_dir = Path(sys.executable).resolve().parent if getattr(sys, 'frozen', False) else Path(__file__).resolve().parent

            template = base_dir / 'mau_hop_dong.docx'
            if not template.exists():
                self.uic9.label_noti.setStyleSheet('color: red')
                self.uic9.label_noti.setText(f'❌ Không tìm thấy template mau_hop_dong.docx tại: {base_dir}')
                return

            out_dir = base_dir / 'hop_dong_out'
            out_dir.mkdir(parents=True, exist_ok=True)
            safe_company = re.sub(r'[^\w\- ]+', '', ten_cty or 'KhachHang').strip().replace(' ', '_')
            default_name = f'Hop_dong_BG_{so_bg}_{safe_company}.docx'
            suggested_path = str(out_dir / default_name)

            save_path, _ = QFileDialog.getSaveFileName(
                self.win_detailcompanyview,
                'Lưu hợp đồng',
                suggested_path,
                'Word Document (*.docx)'
            )
            if not save_path:
                self.uic9.label_noti.setStyleSheet('color: orange')
                self.uic9.label_noti.setText('Đã hủy lưu hợp đồng.')
                return

            if not save_path.lower().endswith('.docx'):
                save_path += '.docx'

            out_path = Path(save_path)

            doc = Document(str(template))

            lead = misc.sql_one(
                'SELECT name, sdt, company, address, mst, email FROM sale_lead WHERE lead_id = (SELECT lead_id FROM ds_bao_gia WHERE so_bg = %s LIMIT 1)',
                (so_bg,)
            )
            ten_lh = lead[0] if lead else ''
            sdt_lh = lead[1] if lead else ''
            lead_company = lead[2] if lead else ''
            lead_address = lead[3] if lead else ''
            lead_mst = lead[4] if lead else ''
            lead_email = lead[5] if lead else ''

            company_row = misc.sql_one(
                'SELECT ten_cong_ty, mst, dia_chi, nguoi_dai_dien, chuc_vu_ndd, sdt_nguoi_dd, email_cong_ty, stk, ngan_hang FROM ds_cong_ty WHERE mst = %s LIMIT 1',
                (mst,)
            ) if mst else None

            if not company_row and ten_cty:
                company_row = misc.sql_one(
                    'SELECT ten_cong_ty, mst, dia_chi, nguoi_dai_dien, chuc_vu_ndd, sdt_nguoi_dd, email_cong_ty, stk, ngan_hang FROM ds_cong_ty WHERE ten_cong_ty = %s LIMIT 1',
                    (ten_cty,)
                )

            cty_ten = company_row[0] if company_row else (lead_company or ten_cty)
            cty_mst = company_row[1] if company_row else (lead_mst or mst)
            cty_dia_chi = company_row[2] if company_row else (lead_address or dia_chi)
            nguoi_dd = company_row[3].strip() if company_row and company_row[3] else ''
            chuc_vu = company_row[4].strip() if company_row and company_row[4] else ''
            cty_sdt_dd = company_row[5].strip() if company_row and company_row[5] else ''
            cty_email = company_row[6].strip() if company_row and company_row[6] else (lead_email or '')
            cty_stk = company_row[7].strip() if company_row and company_row[7] else ''
            cty_bank = company_row[8].strip() if company_row and company_row[8] else ''

            if not company_row:
                self.uic9.label_noti.setStyleSheet('color: red')
                self.uic9.label_noti.setText('❌ Không tìm thấy dữ liệu công ty trong ds_cong_ty theo MST/tên công ty. Vui lòng cập nhật CRM trước khi tạo hợp đồng.')
                return

            if nguoi_dd and (not chuc_vu) and ' - ' in nguoi_dd:
                ten_tach, chuc_vu_tach = nguoi_dd.split(' - ', 1)
                if ten_tach.strip() and chuc_vu_tach.strip():
                    nguoi_dd = ten_tach.strip()
                    chuc_vu = chuc_vu_tach.strip()

            if not nguoi_dd or not chuc_vu:
                self.uic9.label_noti.setStyleSheet('color: red')
                self.uic9.label_noti.setText('❌ Thiếu người đại diện hoặc chức vụ trong ds_cong_ty. Vui lòng cập nhật dữ liệu công ty trước khi tạo hợp đồng.')
                return

            if cty_sdt_dd:
                sdt = cty_sdt_dd

            bg = misc.sql_one('SELECT sotien, noi_dung, sum8, sum10, sum0 FROM ds_bao_gia WHERE so_bg = %s', (so_bg,))
            tong_so = int(bg[0] or 0) if bg else 0
            noi_dung = (bg[1] or '') if bg else ''
            sum8 = int(bg[2] or 0) if bg else 0
            sum10 = int(bg[3] or 0) if bg else 0
            sum0 = int(bg[4] or 0) if bg else 0
            tong_so_fmt = "{:,}".format(tong_so)

            so_hd = f"{so_bg}/HĐKT"
            ngay_ky = datetime.now().strftime('%d/%m/%Y')

            # Tính tổng tiền đúng theo từng dòng báo giá (không suy từ sum8/sum10/sum0)
            tong_hang = 0
            tong_vat = 0
            for line in str(noi_dung).split('@'):
                cols = line.split('|')
                if len(cols) < 7:
                    continue
                try:
                    sl = int(str(cols[4]).replace(',', '').strip() or '0')
                    don_gia = int(str(cols[5]).replace(',', '').strip() or '0')
                    thue = int(str(cols[6]).replace(',', '').strip() or '0')
                except Exception:
                    continue
                thanh_tien = sl * don_gia
                tong_hang += thanh_tien
                if thue == 8:
                    tong_vat += round(thanh_tien * 0.08)
                elif thue == 10:
                    tong_vat += round(thanh_tien * 0.10)
            tong_cong = tong_hang + tong_vat

            def to_vn_words(n: int) -> str:
                if n == 0:
                    return 'Không đồng'
                dv = ['không', 'một', 'hai', 'ba', 'bốn', 'năm', 'sáu', 'bảy', 'tám', 'chín']

                def read_triple(num, full=False):
                    tr = num // 100
                    ch = (num % 100) // 10
                    dvv = num % 10
                    parts = []
                    if full or tr > 0:
                        parts.append(dv[tr] + ' trăm')
                    if ch > 1:
                        parts.append(dv[ch] + ' mươi')
                        if dvv == 1:
                            parts.append('mốt')
                        elif dvv == 5:
                            parts.append('lăm')
                        elif dvv > 0:
                            parts.append(dv[dvv])
                    elif ch == 1:
                        parts.append('mười')
                        if dvv == 5:
                            parts.append('lăm')
                        elif dvv > 0:
                            parts.append(dv[dvv])
                    elif ch == 0 and dvv > 0:
                        if tr > 0 or full:
                            parts.append('lẻ')
                        parts.append(dv[dvv])
                    return ' '.join(parts).strip()

                units = ['', ' nghìn', ' triệu', ' tỷ', ' nghìn tỷ', ' triệu tỷ']
                chunks = []
                x = int(n)
                while x > 0:
                    chunks.append(x % 1000)
                    x //= 1000

                texts = []
                for i in range(len(chunks) - 1, -1, -1):
                    c = chunks[i]
                    if c == 0:
                        continue
                    full = i < len(chunks) - 1 and chunks[i + 1] != 0
                    part = read_triple(c, full=full)
                    texts.append(part + units[i])

                sentence = ' '.join(texts).strip()
                sentence = re.sub(r'\s+', ' ', sentence)
                return sentence[:1].upper() + sentence[1:] + ' đồng'

            values = {
                # placeholders trong template hiện tại
                '{so_hd}': so_hd,
                '{contract-date}': ngay_ky,
                '{company_name}': cty_ten,
                '{address}': cty_dia_chi,
                '{email_cty}': cty_email,
                '{stk}': cty_stk,
                '{bank}': cty_bank,
                '{mst}': cty_mst,
                '{dai_dien}': nguoi_dd,
                '{chuc_vu}': chuc_vu,
                '{tien-bang-so}': '{:,}'.format(tong_cong),
                '{tien-bang-chu}': to_vn_words(tong_cong),
                '{tong_tien_hang}': '{:,}'.format(tong_hang),
                '{tong_vat}': '{:,}'.format(tong_vat),
                '{tong_cong}': '{:,}'.format(tong_cong),

                # giữ thêm các key cũ/phòng hờ
                '{SO_BG}': str(so_bg),
                '{NGAY}': ngay_ky,
                '{TEN_CONG_TY}': cty_ten,
                '{MST}': cty_mst,
                '{DIA_CHI}': cty_dia_chi,
                '{NGUOI_DAI_DIEN}': nguoi_dd,
                '{CHUC_VU}': chuc_vu,
                '{SDT}': sdt_lh or sdt,
                '{TEN_LIEN_HE}': ten_lh,
                '{SDT_LIEN_HE}': sdt_lh,
            }

            def replace_in_paragraph(p):
                text = p.text
                new_text = text
                for k, v in values.items():
                    new_text = new_text.replace(k, v or '')
                if new_text != text:
                    p.text = new_text

            for p in doc.paragraphs:
                replace_in_paragraph(p)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p)

            # Đổ danh mục hàng hóa vào bảng placeholder {{ITEM_TABLE}}
            goods = []
            for line in str(noi_dung).split('@'):
                cols = line.split('|')
                if len(cols) < 7:
                    continue
                try:
                    goods.append({
                        'ten': cols[0],
                        'model': cols[1],
                        'dvt': cols[3],
                        'sl': int(str(cols[4]).replace(',', '').strip() or '0'),
                        'don_gia': int(str(cols[5]).replace(',', '').strip() or '0'),
                        'thue': int(str(cols[6]).replace(',', '').strip() or '0'),
                    })
                except Exception:
                    continue

            marker_table = None
            marker_row_idx = -1
            for tb in doc.tables:
                found = False
                for ri, rw in enumerate(tb.rows):
                    for c in rw.cells:
                        if '{{ITEM_TABLE}}' in (c.text or ''):
                            marker_table = tb
                            marker_row_idx = ri
                            found = True
                            break
                    if found:
                        break
                if found:
                    break

            if marker_table is not None and marker_row_idx >= 0:
                # Xóa dòng marker
                marker_table._tbl.remove(marker_table.rows[marker_row_idx]._tr)

                col_count = len(marker_table.columns)

                for idx, g in enumerate(goods, start=1):
                    rw = marker_table.add_row().cells
                    thanh_tien = int(g['sl']) * int(g['don_gia'])

                    # Cấu trúc theo template không có cột thuế (7 cột)
                    if col_count >= 7:
                        rw[0].text = str(idx)
                        rw[1].text = str(g['ten'])
                        rw[2].text = str(g['model'])
                        rw[3].text = str(g['dvt'])
                        rw[4].text = str(g['sl'])
                        rw[5].text = '{:,}'.format(int(g['don_gia']))
                        if col_count == 7:
                            rw[6].text = '{:,}'.format(thanh_tien)
                        else:
                            # nếu template còn cột thuế
                            rw[6].text = str(g['thue'])
                            rw[7].text = '{:,}'.format(thanh_tien)

            doc.save(str(out_path))
            self.uic9.label_noti.setStyleSheet('color: blue')
            self.uic9.label_noti.setText(f'✅ Đã tạo hợp đồng: {out_path.name}')

        except Exception as e:
            self.uic9.label_noti.setStyleSheet('color: red')
            self.uic9.label_noti.setText(f'❌ Lỗi tạo hợp đồng: {e}')

    def add_company(self):
        self.win_addcompanyview = QMainWindow()
        self.uic9 = Ui_ViewDetailCompany()
        self.uic9.setupUi(self.win_addcompanyview)
        apply_ui_v2(self.win_addcompanyview)
        self.win_addcompanyview.show()

        self.uic9.label_noti.setText("Thêm thông tin khách hàng mới vào CSDL")
        self.uic9.tabWidget.setTabEnabled(1, False)
        self.uic9.tabWidget.setTabEnabled(2, False)
        self.uic9.but_tao_lead.hide()
        self.uic9.but_tao_moi.clicked.connect(lambda: Crm.update_tt_cong_ty(self))
        self.uic9.but_update.clicked.connect(lambda: Crm.update_tt_cong_ty(self))

    def search_company(self):
        current_tab_index = self.uic8.tabWidget.currentIndex()
        if current_tab_index == self.uic8.tabWidget.indexOf(self.uic8.tab):
            self.uic8.tableWidget.clear()
            self.search_company_text = self.uic8.textEdit.toPlainText()

            # Tối ưu: dùng 1 câu query UNION để giảm round-trip DB + tránh merge Python O(n^2)
            search_text = f"%{self.search_company_text}%"
            query = """
                SELECT * FROM ds_cong_ty WHERE mst LIKE %s
                UNION
                SELECT * FROM ds_cong_ty WHERE ten_cong_ty LIKE %s
                UNION
                SELECT * FROM ds_cong_ty WHERE dia_chi LIKE %s
                LIMIT 1000
            """
            result = misc.sql_all(query, (search_text, search_text, search_text))
            if result:
                self.uic8.tableWidget.setColumnCount(5)
                self.uic8.tableWidget.setRowCount(len(result))
                self.uic8.tableWidget.setHorizontalHeaderLabels(
                    ['Mã KH', 'Mã số thuế', 'Tên công ty', 'Địa chỉ', 'Thao tác'])
                self.uic8.tableWidget.setColumnWidth(0, 40)
                self.uic8.tableWidget.setColumnWidth(1, 80)
                self.uic8.tableWidget.setColumnWidth(2, 200)
                self.uic8.tableWidget.setColumnWidth(3, 220)
                self.uic8.tableWidget.setColumnWidth(4, 50)

                for row in range(len(result)):
                    self.uic8.tableWidget.setItem(row, 0, QTableWidgetItem(str(result[row][0])))
                    self.uic8.tableWidget.setItem(row, 1, QTableWidgetItem(result[row][2]))
                    self.uic8.tableWidget.setItem(row, 2, QTableWidgetItem(result[row][1]))
                    self.uic8.tableWidget.setItem(row, 3, QTableWidgetItem(result[row][4]))

                    # self.uic8.tableWidget.resizeRowToContents(row)

                    but1 = QPushButton('Xem')
                    but1.clicked.connect(lambda _, r=row: Crm.view_detail_company(self, result[r][2]))
                    self.uic8.tableWidget.setCellWidget(row, 4, but1)

                    # Get the current height of the row
                    current_row_height = self.uic8.tableWidget.rowHeight(row)

                    # Check if current row height exceeds maximum row height
                    # if current_row_height > 65:
                    self.uic8.tableWidget.setRowHeight(row, 70)

                self.uic8.tableWidget.repaint()
            else:
                self.uic8.tableWidget.clear()
                print('Lỗi không hiển thị được')
        else:
            self.uic8.tableWidget_2.clear()
            self.search_company_text = self.uic8.textEdit.toPlainText()

            # Tối ưu: query UNION, giảm tải DB và tránh dedupe Python
            search_text = f"%{self.search_company_text}%"
            query = """
                SELECT * FROM ds_ca_nhan WHERE dien_thoai LIKE %s
                UNION
                SELECT * FROM ds_ca_nhan WHERE ten LIKE %s
                UNION
                SELECT * FROM ds_ca_nhan WHERE ten_cong_ty LIKE %s
                LIMIT 1000
            """
            result1 = misc.sql_all(query, (search_text, search_text, search_text))
            if result1:
                self.uic8.tableWidget_2.setColumnCount(4)
                self.uic8.tableWidget_2.setRowCount(len(result1))
                self.uic8.tableWidget_2.setHorizontalHeaderLabels(
                    ['Điện thoại', 'Tên khách hàng', 'Tên công ty', 'Thao tác'])
                self.uic8.tableWidget_2.setColumnWidth(0, 100)
                self.uic8.tableWidget_2.setColumnWidth(1, 200)
                self.uic8.tableWidget_2.setColumnWidth(2, 200)
                self.uic8.tableWidget_2.setColumnWidth(3, 70)

                for row in range(len(result1)):
                    self.uic8.tableWidget_2.setItem(row, 0, QTableWidgetItem(str(result1[row][1])))

                    self.uic8.tableWidget_2.setItem(row, 1, QTableWidgetItem(result1[row][0]))
                    self.uic8.tableWidget_2.setItem(row, 2, QTableWidgetItem(result1[row][6]))

                    self.uic8.tableWidget_2.resizeRowToContents(row)

                    but1 = QPushButton('Xem')
                    but1.clicked.connect(lambda _, r=row: Crm.view_detail_canhan(self, r))

                    self.uic8.tableWidget_2.setCellWidget(row, 3, but1)

                    self.uic8.tableWidget_2.setRowHeight(row, 65)

                self.uic8.tableWidget_2.repaint()
            else:
                print('Không tìm thấy kết quả!!!!')
                self.uic8.tableWidget_2.clear()
