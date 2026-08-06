# -*- coding: utf-8 -*-
"""
Module xuất 4 mẫu PC01-PC04 theo NĐ 105/2025/NĐ-CP về PCCC&CNCH.

- PC01: Phiếu thông tin của cơ sở (khai lần đầu)
- PC02: Biên bản tự kiểm tra PCCC (định kỳ)
- PC03: Biên bản kiểm tra của CA (có cơ quan NN)
- PC04: Báo cáo kết quả thực hiện công tác PCCC

Dữ liệu input: dict từ khao_sat_data.thu_thap() — schema v5.
"""
import json
import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Mapping option → text hiển thị
HINH_THUC_LBL = {"so_huu": "Sở hữu", "thue": "Thuê", "khac": "Khác"}
TRANG_THAI_LBL = {"thiet_ke": "Đang thiết kế",
                  "thi_cong": "Đang thi công",
                  "van_hanh": "Đã đi vào vận hành"}
NGUON_NUOC_LBL = {"duong_ong": "Đường ống đô thị",
                  "be_chua": "Bể chứa riêng",
                  "song_ho": "Sông / hồ / ao"}
NGUON_DIEN_LBL = {"1_nguon": "1 nguồn", "2_nguon": "2 nguồn độc lập"}
THANH_PHAN_KT_LBL = {"nha_nuoc": "Nhà nước",
                     "tap_the": "Tập thể",
                     "tu_nhan": "Tư nhân",
                     "von_nuoc_ngoai": "Có vốn đầu tư nước ngoài"}
TAI_LIEU_LBL = {
    "gp_xay_dung": "Giấy phép xây dựng",
    "gcn_qsdd": "GCN quyền sử dụng đất",
    "bv_thiet_ke_pccc": "Bản vẽ thiết kế PCCC",
    "bv_hoan_cong": "Bản vẽ hoàn công",
    "bb_nghiem_thu_pccc": "Biên bản nghiệm thu PCCC",
    "qd_tham_duyet": "Quyết định thẩm duyệt PCCC trước đó",
    "so_theo_doi_pccc": "Sổ theo dõi PCCC",
    "bb_kiem_tra_ca": "Biên bản kiểm tra của CA PCCC",
    "hd_su_dung": "Hướng dẫn sử dụng thiết bị PCCC",
    "bh_chay_no": "Hợp đồng bảo hiểm cháy nổ",
}


# =====================================================================
# Helpers
# =====================================================================
def _add_p(doc, text="", bold=False, size=11, align=None,
           color=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _add_field(doc, label, value, dots="…" * 30):
    """Thêm 1 hàng dạng: 'Label: ...........value............'"""
    p = doc.add_paragraph()
    r = p.add_run(label + ": ")
    r.bold = False
    r.font.size = Pt(11)
    if value and str(value).strip():
        r2 = p.add_run(str(value))
        r2.bold = True
        r2.font.size = Pt(11)
    else:
        r3 = p.add_run(dots)
        r3.font.size = Pt(11)
        r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


def _add_checkbox(doc, label, checked=False):
    """Thêm 1 hàng dạng: '- Label: ☐ / ☑'"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("- " + label + ": ")
    r.font.size = Pt(11)
    r2 = p.add_run("☑" if checked else "☐")
    r2.bold = True
    r2.font.size = Pt(12)
    return p


def _add_section_header(doc, title, level=1):
    """Header section."""
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    if level == 1:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    else:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def _add_header_cong_hoa(doc, ten_co_so=""):
    """Header: CỘNG HOÀ XHCN VIỆT NAM + Tên cơ sở (2 cột)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    cell_l, cell_r = table.rows[0].cells
    # Left: tên cơ sở
    pl = cell_l.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = pl.add_run((ten_co_so or "TÊN CƠ SỞ").upper())
    rl.bold = True
    rl.font.size = Pt(11)
    pl2 = cell_l.add_paragraph()
    pl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl2 = pl2.add_run("–––––––––––")
    rl2.font.size = Pt(10)
    # Right: Cộng hòa
    pr = cell_r.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pr.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    rr.bold = True
    rr.font.size = Pt(11)
    pr2 = cell_r.add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr2 = pr2.add_run("Độc lập - Tự do - Hạnh phúc")
    rr2.bold = True
    rr2.italic = True
    rr2.font.size = Pt(11)
    pr3 = cell_r.add_paragraph()
    pr3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr3.add_run("–––––––––––").font.size = Pt(10)
    doc.add_paragraph()


def _add_signature(doc, vai_tro="NGƯỜI ĐỨNG ĐẦU CƠ SỞ",
                   noi_ngay=None, name=""):
    """Khu ký tên — căn phải."""
    doc.add_paragraph()
    noi_ngay = noi_ngay or _today_str()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(noi_ngay)
    r.italic = True
    r.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(vai_tro)
    r2.bold = True
    r2.font.size = Pt(11)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run("(Ký, ghi rõ họ tên)")
    r3.italic = True
    r3.font.size = Pt(10)
    for _ in range(3):
        doc.add_paragraph()
    if name:
        pn = doc.add_paragraph()
        pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rn = pn.add_run(name)
        rn.bold = True
        rn.font.size = Pt(11)


def _today_str():
    """Trả: '…, ngày DD tháng MM năm YYYY'"""
    d = datetime.date.today()
    return f"…, ngày {d.day:02d} tháng {d.month:02d} năm {d.year}"


def _v(data, key, default=""):
    """Get value, fallback default."""
    v = data.get(key)
    if v is None or v == "":
        return default
    return v


def _bool_chk(v):
    """Convert 0/1/'1'/True → bool."""
    if v in (1, "1", True, "True", "true"):
        return True
    return False


# =====================================================================
# PC01 — Phiếu thông tin của cơ sở (khai lần đầu)
# =====================================================================
def xuat_pc01(data: dict, file_out: str, nguoi_dung_dau="") -> str:
    """Xuất Mẫu PC01 — Phiếu thông tin của cơ sở."""
    doc = Document()

    # === Header ===
    _add_p(doc, "Mẫu số PC01", italic=True, size=10, align="right")
    _add_p(doc, "PHIẾU THÔNG TIN CỦA CƠ SỞ",
           bold=True, size=15, align="center",
           color=(0x0F, 0x76, 0x6E))
    _add_p(doc, "(KHAI BÁO LẦN ĐẦU)",
           italic=True, size=11, align="center")
    doc.add_paragraph()

    # === I. THÔNG TIN CHUNG ===
    _add_section_header(doc, "I. THÔNG TIN CHUNG VỀ CƠ SỞ")

    _add_field(doc, "1. Tên cơ sở",
               _v(data, "kh_cty") or _v(data, "ten_cong_trinh"))
    _add_field(doc, "2. Địa chỉ", _v(data, "dia_chi"))
    _add_field(doc, "3. Ngành nghề, lĩnh vực hoạt động",
               _v(data, "nganh_nghe"))
    _add_field(doc, "4. Năm đưa vào hoạt động",
               _v(data, "nam_hoat_dong"))
    _add_field(doc, "5. Tên cơ quan/tổ chức/cá nhân trực tiếp quản lý",
               _v(data, "nguoi_quan_ly"))
    _add_field(doc,
               "6. Họ tên người đứng đầu cơ sở/đại diện pháp luật",
               _v(data, "dai_dien_pl_ten"))
    _add_field(doc, "   Số điện thoại", _v(data, "dai_dien_pl_sdt"))
    _add_field(doc, "7. Tên cơ quan/tổ chức cấp trên (nếu có)",
               _v(data, "co_quan_cap_tren"))

    # 8. Thành phần kinh tế (checkbox)
    p = doc.add_paragraph()
    p.add_run("8. Thuộc thành phần kinh tế:").bold = False
    tpkt = _v(data, "thanh_phan_kt")
    for k, lbl in [("nha_nuoc", "Nhà nước"),
                   ("tap_the", "Tập thể"),
                   ("tu_nhan", "Tư nhân"),
                   ("von_nuoc_ngoai", "Có vốn đầu tư nước ngoài")]:
        _add_checkbox(doc, lbl, checked=(tpkt == k))

    _add_checkbox(doc,
                  "9. Thuộc danh mục cơ sở có nguy hiểm về cháy, nổ",
                  checked=_bool_chk(_v(data, "thuoc_dm_nguyhiem")))
    _add_checkbox(doc,
                  "10. Thuộc danh mục dự án phải thẩm duyệt PCCC",
                  checked=_bool_chk(_v(data, "thuoc_dm_thamduyet")))

    doc.add_paragraph()

    # === II. THÔNG TIN LIÊN QUAN PCCC ===
    _add_section_header(doc,
                        "II. THÔNG TIN LIÊN QUAN PCCC, CNCH")

    # 1. Pháp lý PCCC
    _add_section_header(doc,
                        "1. Văn bản pháp lý về PCCC (nếu có):", level=2)
    _add_field(doc,
               "   - Văn bản thẩm duyệt thiết kế PCCC (số)",
               _v(data, "vb_thamduyet_so"))
    _add_field(doc, "     Ngày ban hành",
               _v(data, "vb_thamduyet_ngay"))
    _add_field(doc, "     Cơ quan ban hành",
               _v(data, "vb_thamduyet_cq"))
    _add_field(doc, "   - Văn bản chấp thuận nghiệm thu PCCC (số)",
               _v(data, "vb_nghiemthu_so"))
    _add_field(doc, "     Ngày ban hành",
               _v(data, "vb_nghiemthu_ngay"))
    _add_field(doc, "     Cơ quan ban hành",
               _v(data, "vb_nghiemthu_cq"))

    # 2. Quy mô + đặc điểm
    _add_section_header(doc,
                        "2. Quy mô + tính chất nguy hiểm cháy, nổ:",
                        level=2)
    _add_field(doc, "   - Tổng DT sàn sử dụng (m²)",
               _v(data, "dt_san_tong"))
    _add_field(doc, "     DT xây dựng (m²)",
               _v(data, "dt_xay_dung"))
    _add_field(doc, "     Chiều cao PCCC (m)",
               _v(data, "cao_pccc"))
    _add_field(doc, "     Số tầng nổi", _v(data, "so_tang_noi"))
    _add_field(doc, "     Số tầng hầm", _v(data, "so_tang_ham"))
    _add_field(doc, "     Bậc chịu lửa",
               _v(data, "bac_chiu_lua"))
    _add_field(doc, "     Cấp nguy hiểm cháy KC",
               _v(data, "cap_nhc"))
    _add_field(doc, "     Hạng nguy hiểm cháy",
               _v(data, "hang_nguy_hiem"))
    _add_field(doc, "     Số người sử dụng",
               _v(data, "so_nguoi_du_kien"))

    # Bảng khối nhà nếu có
    try:
        kn_list = json.loads(_v(data, "khoi_nha_json") or "[]")
    except Exception:
        kn_list = []
    if kn_list:
        _add_p(doc, "   - Các khối nhà trong cơ sở:",
               size=11)
        t = doc.add_table(rows=1 + len(kn_list), cols=7)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(
                ["Tên khối", "DTXD (m²)", "Tầng nổi", "Tầng hầm",
                 "Bậc CL", "Công năng", "Lối thoát"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for i, kn in enumerate(kn_list):
            row = t.rows[1 + i].cells
            row[0].text = kn.get("ten_khoi", "")
            row[1].text = str(kn.get("dt_xay_dung", ""))
            row[2].text = str(kn.get("so_tang_noi", ""))
            row[3].text = str(kn.get("so_tang_ham", ""))
            row[4].text = kn.get("bac_chiu_lua", "")
            row[5].text = kn.get("cong_nang", "")
            row[6].text = str(kn.get("so_loi_thoat", ""))

    # Khu vực ngoài nhà
    try:
        kv_list = json.loads(_v(data, "khu_vuc_json") or "[]")
    except Exception:
        kv_list = []
    if kv_list:
        _add_p(doc, "   - Các khu vực NGOÀI nhà:", size=11)
        t = doc.add_table(rows=1 + len(kv_list), cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(
                ["Tên khu vực", "DT (m²)", "Dây chuyền CN / vật tư"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for i, kv in enumerate(kv_list):
            row = t.rows[1 + i].cells
            row[0].text = kv.get("ten_khu_vuc", "")
            row[1].text = str(kv.get("dt_su_dung", ""))
            row[2].text = kv.get("day_chuyen_cn", "")

    # 3. Giao thông
    _add_section_header(doc,
                        "3. Giao thông phục vụ chữa cháy, CNCH:",
                        level=2)
    _add_field(doc, "   - Chiều rộng đường giao thông (m)",
               _v(data, "dgt_rong"))
    _add_field(doc, "     Chiều cao thông thủy (m)",
               _v(data, "dgt_cao"))
    _add_field(doc, "   - Vị trí bãi đỗ xe CC",
               _v(data, "bai_do_xe_cc"))

    # 4. Nguồn nước CC
    _add_section_header(doc, "4. Nguồn nước phục vụ chữa cháy:",
                        level=2)
    _add_field(doc, "   - Số bể CC", _v(data, "so_be_cc"))
    _add_field(doc, "     Khối tích bể (m³)",
               _v(data, "khoi_tich_be"))
    _add_field(doc, "     Vị trí bể + khả năng lấy nước",
               _v(data, "vi_tri_be"))
    _add_field(doc, "   - Số trụ cấp nước CC",
               _v(data, "so_tru_cc"))
    _add_field(doc, "     Vị trí trụ", _v(data, "vi_tri_tru"))
    nn_list = []
    try:
        nn_list = json.loads(_v(data, "nguon_nuoc_json") or "[]")
    except Exception:
        pass
    nn_text = ", ".join(NGUON_NUOC_LBL.get(k, k) for k in nn_list)
    _add_field(doc, "   - Loại nguồn nước", nn_text)

    # 5. Hệ thống PCCC + phương tiện
    _add_section_header(
        doc,
        "5. Hệ thống PCCC + phương tiện CC, CNCH của cơ sở:",
        level=2)
    try:
        ht_list = json.loads(_v(data, "he_thong_sn_json") or "[]")
    except Exception:
        ht_list = []
    if ht_list:
        t = doc.add_table(rows=1 + len(ht_list), cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Hệ thống", "Tình trạng", "Hãng/Model"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for i, ht in enumerate(ht_list):
            row = t.rows[1 + i].cells
            row[0].text = ht.get("ten", "")
            row[1].text = ht.get("tinh_trang", "")
            row[2].text = ht.get("hang", "")
    _add_field(doc, "   - Hệ thống điện PCCC: nguồn lưới + dự phòng",
               f"Máy phát {_v(data, 'mp_cong_suat_kva')} kVA, chạy "
               f"{_v(data, 'mp_thoi_gian_chay_h')} giờ"
               if _bool_chk(_v(data, "co_may_phat")) else "")
    _add_field(doc, "   - Đã lắp truyền tin báo cháy",
               "Có" if _bool_chk(_v(data, "truyen_tin_da_lap"))
               else "Chưa")
    _add_field(doc, "   - Phương tiện CC cơ giới",
               _v(data, "phuong_tien_cc_text"))

    # 6. Tổ chức lực lượng tại chỗ
    _add_section_header(doc, "6. Tổ chức lực lượng tại chỗ:",
                        level=2)
    _add_field(doc, "   - Tổng số đội viên PCCC cơ sở",
               _v(data, "doi_tong_doi_vien"))
    _add_field(doc, "   - Họ tên đội trưởng",
               _v(data, "doi_truong_ten"))
    _add_field(doc, "     SĐT", _v(data, "doi_truong_sdt"))
    _add_field(doc, "   - Tổng người được phân công PCCC",
               _v(data, "so_nguoi_pccc"))

    # === Ký tên ===
    _add_signature(doc, vai_tro="NGƯỜI ĐỨNG ĐẦU CƠ SỞ",
                   name=nguoi_dung_dau)

    doc.save(file_out)
    return file_out


# =====================================================================
# PC02 — Biên bản tự kiểm tra PCCC (định kỳ)
# =====================================================================
def xuat_pc02(data: dict, file_out: str,
              nguoi_kiem_tra="", chuc_vu="") -> str:
    doc = Document()
    _add_p(doc, "Mẫu số PC02", italic=True, size=10, align="right")
    _add_p(doc, "BIÊN BẢN TỰ KIỂM TRA",
           bold=True, size=15, align="center",
           color=(0x0F, 0x76, 0x6E))
    _add_p(doc, "Về phòng cháy, chữa cháy",
           italic=True, size=11, align="center")
    doc.add_paragraph()

    now = datetime.datetime.now()
    _add_p(doc,
           f"Hồi {now.hour:02d} giờ {now.minute:02d} ngày "
           f"{now.day:02d} tháng {now.month:02d} năm {now.year}, "
           f"chúng tôi gồm:",
           size=11, align="justify")
    _add_field(doc, "- Ông/bà", nguoi_kiem_tra)
    _add_field(doc, "  Chức vụ", chuc_vu)
    doc.add_paragraph()

    _add_p(doc, "Đã tiến hành kiểm tra đối với:", size=11)
    _add_field(doc, "Cơ sở",
               _v(data, "kh_cty") or _v(data, "ten_cong_trinh"))
    _add_field(doc, "Địa chỉ", _v(data, "dia_chi"))
    doc.add_paragraph()

    _add_section_header(doc,
                        "1. Nội dung và kết quả kiểm tra:")

    # 1a. Phương tiện CC, hệ thống điện, nguồn nước
    _add_p(doc,
           "- Việc trang bị phương tiện, hệ thống PCCC, CNCH, "
           "hệ thống điện PCCC, nguồn nước CC:",
           bold=True, size=11)
    _add_p(doc,
           "  + Chủng loại, số lượng, vị trí phương tiện CC, CNCH:",
           size=11)
    try:
        ht_list = json.loads(_v(data, "he_thong_sn_json") or "[]")
    except Exception:
        ht_list = []
    if ht_list:
        t = doc.add_table(rows=1 + len(ht_list), cols=4)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(
                ["STT", "Hệ thống", "Tình trạng", "Hãng/Model"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for i, ht in enumerate(ht_list):
            row = t.rows[1 + i].cells
            row[0].text = str(i + 1)
            row[1].text = ht.get("ten", "")
            row[2].text = ht.get("tinh_trang", "")
            row[3].text = ht.get("hang", "")

    _add_p(doc,
           "  + Hệ thống điện PCCC (nguồn lưới + dự phòng):",
           size=11)
    _add_field(doc, "     Có máy phát dự phòng",
               "Có" if _bool_chk(_v(data, "co_may_phat")) else "Không")
    _add_field(doc, "     Công suất máy phát (kVA)",
               _v(data, "mp_cong_suat_kva"))
    _add_field(doc, "     Thời gian chạy (giờ)",
               _v(data, "mp_thoi_gian_chay_h"))

    _add_p(doc, "  + Nguồn nước CC:", size=11)
    _add_field(doc, "     Số bể CC", _v(data, "so_be_cc"))
    _add_field(doc, "     Khối tích bể (m³)",
               _v(data, "khoi_tich_be"))
    _add_field(doc, "     Số trụ cấp nước CC",
               _v(data, "so_tru_cc"))

    # 1b. Truyền tin báo cháy
    _add_p(doc,
           "- Lắp đặt thiết bị truyền tin báo cháy kết nối CSDL PCCC:",
           bold=True, size=11)
    _add_field(doc, "  Đã lắp truyền tin báo cháy",
               "Có" if _bool_chk(_v(data, "truyen_tin_da_lap"))
               else "Chưa lắp")

    # 1c. Khoảng cách + giao thông CC
    _add_p(doc,
           "- Khoảng cách PCCC, đường, bãi đỗ phục vụ CC, CNCH:",
           bold=True, size=11)
    _add_field(doc, "  Khoảng cách tới CT kế bên (m)",
               _v(data, "kc_ct_ke_ben"))
    _add_field(doc, "  Đường giao thông CC — Rộng (m)",
               _v(data, "dgt_rong"))
    _add_field(doc, "    Cao thông thủy (m)", _v(data, "dgt_cao"))
    _add_field(doc, "    Vị trí bãi đỗ xe CC",
               _v(data, "bai_do_xe_cc"))

    # 1d. Thoát nạn
    _add_p(doc,
           "- Giải pháp thoát nạn, ngăn cháy, chống cháy lan, "
           "chống khói:",
           bold=True, size=11)
    _add_p(doc,
           "  (Kiểm tra thực tế lối thoát, đường thoát, chiều rộng "
           "cửa thoát nạn).",
           italic=True, size=10)

    # 1e. Lịch sử PCCC + sơ hở
    _add_p(doc,
           "- Sơ hở, thiếu sót đã được phát hiện trước đây + "
           "việc khắc phục:",
           bold=True, size=11)
    if _v(data, "lich_su_pccc"):
        _add_p(doc, _v(data, "lich_su_pccc"),
               size=11, align="justify")
    else:
        _add_p(doc, "(Không có)", italic=True, size=10)

    doc.add_paragraph()
    _add_section_header(doc, "2. Kiến nghị:")
    _add_p(doc,
           "(Ghi yêu cầu cụ thể thời hạn khắc phục các sơ hở, "
           "thiếu sót, vi phạm PCCC, CNCH; kiến nghị, đề xuất "
           "với người có thẩm quyền)",
           italic=True, size=10)
    for _ in range(5):
        _add_p(doc, "…" * 80, size=11)

    _add_p(doc,
           f"Việc kiểm tra kết thúc hồi {now.hour:02d}:{now.minute:02d} "
           f"ngày {now.day:02d}/{now.month:02d}/{now.year}.",
           size=11, italic=True)

    _add_signature(doc, vai_tro="NGƯỜI KIỂM TRA",
                   name=nguoi_kiem_tra)

    doc.save(file_out)
    return file_out


# =====================================================================
# PC03 — Biên bản kiểm tra (có cơ quan NN)
# =====================================================================
def xuat_pc03(data: dict, file_out: str) -> str:
    doc = Document()
    _add_p(doc, "Mẫu số PC03", italic=True, size=10, align="right")
    _add_p(doc, "BIÊN BẢN KIỂM TRA",
           bold=True, size=15, align="center",
           color=(0x0F, 0x76, 0x6E))
    _add_p(doc, "Về phòng cháy, chữa cháy",
           italic=True, size=11, align="center")
    doc.add_paragraph()

    now = datetime.datetime.now()
    _add_p(doc,
           f"Hồi {now.hour:02d} giờ {now.minute:02d} ngày "
           f"{now.day:02d} tháng {now.month:02d} năm {now.year}, tại:",
           size=11, align="justify")
    _add_field(doc, "Cơ sở",
               _v(data, "kh_cty") or _v(data, "ten_cong_trinh"))
    _add_field(doc, "Địa chỉ", _v(data, "dia_chi"))
    doc.add_paragraph()

    _add_p(doc, "Chúng tôi gồm:", size=11, bold=True)
    _add_p(doc, "Đại diện cơ quan kiểm tra:", size=11)
    _add_field(doc, "- Ông/bà", "")
    _add_field(doc, "  Chức vụ", "")
    _add_field(doc, "- Ông/bà", "")
    _add_field(doc, "  Chức vụ", "")
    doc.add_paragraph()

    _add_p(doc, "Đại diện cơ sở được kiểm tra:", size=11)
    _add_field(doc, "- Ông/bà", _v(data, "dai_dien_pl_ten"))
    _add_field(doc, "  SĐT", _v(data, "dai_dien_pl_sdt"))
    _add_field(doc, "  Người tiếp đón", _v(data, "td_ten"))
    _add_field(doc, "  Chức vụ", _v(data, "td_chuc_vu"))
    _add_field(doc, "  SĐT", _v(data, "td_sdt"))
    doc.add_paragraph()

    _add_p(doc,
           "Đã tiến hành kiểm tra về phòng cháy chữa cháy đối với "
           "cơ sở trên.",
           size=11)
    doc.add_paragraph()

    _add_section_header(doc, "1. Nội dung và kết quả kiểm tra:")
    _add_p(doc, "a) Nội dung trình bày của đại diện cơ sở:",
           bold=True, size=11)
    for _ in range(4):
        _add_p(doc, "…" * 80, size=11)

    _add_p(doc, "b) Kết quả kiểm tra:", bold=True, size=11)
    # Tóm tắt từ data
    _add_field(doc, "- Tên cơ sở",
               _v(data, "kh_cty") or _v(data, "ten_cong_trinh"))
    _add_field(doc, "- Năm hoạt động", _v(data, "nam_hoat_dong"))
    _add_field(doc, "- Tổng DT sàn (m²)", _v(data, "dt_san_tong"))
    _add_field(doc, "- Số tầng nổi/hầm",
               f"{_v(data, 'so_tang_noi')}/{_v(data, 'so_tang_ham')}")
    _add_field(doc, "- Bậc chịu lửa", _v(data, "bac_chiu_lua"))
    _add_field(doc, "- VB Thẩm duyệt PCCC",
               f"{_v(data, 'vb_thamduyet_so')} ngày "
               f"{_v(data, 'vb_thamduyet_ngay')} - "
               f"{_v(data, 'vb_thamduyet_cq')}")
    _add_field(doc, "- VB Nghiệm thu PCCC",
               f"{_v(data, 'vb_nghiemthu_so')} ngày "
               f"{_v(data, 'vb_nghiemthu_ngay')} - "
               f"{_v(data, 'vb_nghiemthu_cq')}")

    _add_p(doc, "c) Sơ hở, thiếu sót phát hiện:", bold=True, size=11)
    if _v(data, "lich_su_pccc"):
        _add_p(doc, _v(data, "lich_su_pccc"),
               size=11, align="justify")
    else:
        for _ in range(4):
            _add_p(doc, "…" * 80, size=11)

    doc.add_paragraph()
    _add_section_header(doc, "2. Kiến nghị:")
    _add_p(doc,
           "(Ghi yêu cầu cụ thể thời hạn khắc phục các sơ hở, "
           "thiếu sót, vi phạm PCCC, CNCH; hướng dẫn về PCCC, "
           "CNCH (nếu có))",
           italic=True, size=10)
    for _ in range(5):
        _add_p(doc, "…" * 80, size=11)

    _add_p(doc,
           f"Biên bản được lập xong hồi {now.hour:02d}:{now.minute:02d} "
           f"ngày {now.day:02d}/{now.month:02d}/{now.year}, gồm … trang, "
           f"được lập thành … bản, mỗi bên giữ 01 bản, "
           f"đã đọc lại cho mọi người cùng nghe, công nhận đúng và "
           f"nhất trí ký tên dưới đây.",
           size=11, align="justify", italic=True)

    # Bảng ký 3 cột
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=3)
    cells = t.rows[0].cells
    for i, vt in enumerate([
            "ĐẠI DIỆN CƠ SỞ\n(Ký, ghi rõ họ tên,\nđóng dấu nếu có)",
            "ĐẠI DIỆN CÓ LIÊN QUAN\n(nếu có)",
            "ĐOÀN KIỂM TRA\n(Ký, ghi rõ họ tên, chức vụ)"]):
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cells[i].paragraphs[0].add_run(vt)
        r.bold = True
        r.font.size = Pt(11)

    doc.save(file_out)
    return file_out


# =====================================================================
# PC04 — Báo cáo kết quả thực hiện công tác PCCC
# =====================================================================
def xuat_pc04(data: dict, file_out: str,
              nguoi_dung_dau="") -> str:
    doc = Document()
    _add_p(doc, "Mẫu số PC04", italic=True, size=10, align="right")

    # Header: Tên cơ sở + Cộng hòa
    _add_header_cong_hoa(doc,
                         _v(data, "kh_cty") or
                         _v(data, "ten_cong_trinh"))

    _add_p(doc, "BÁO CÁO",
           bold=True, size=15, align="center",
           color=(0x0F, 0x76, 0x6E))
    _add_p(doc, "Kết quả thực hiện công tác phòng cháy, chữa cháy, "
                "cứu nạn, cứu hộ",
           bold=True, size=11, align="center")
    doc.add_paragraph()

    _add_p(doc, "Kính gửi: (Ủy ban nhân dân cấp xã / Cơ quan Công an "
                "/ Cơ quan chuyên môn về xây dựng trực tiếp quản lý "
                "cơ sở)",
           italic=True, size=10)
    doc.add_paragraph()

    _add_p(doc,
           f"({_v(data, 'kh_cty') or _v(data, 'ten_cong_trinh')})"
           " báo cáo kết quả thực hiện công tác PCCC, CNCH của cơ sở "
           "như sau:",
           size=11, align="justify")
    doc.add_paragraph()

    _add_section_header(doc, "1. Kết quả thực hiện:")

    _add_p(doc,
           "a) Trách nhiệm về PCCC, CNCH của người đứng đầu cơ sở "
           "(theo khoản 3 Điều 8 Luật PCCC):",
           bold=True, size=11)
    _add_field(doc, "   Người đứng đầu cơ sở",
               _v(data, "dai_dien_pl_ten"))
    _add_field(doc, "   SĐT", _v(data, "dai_dien_pl_sdt"))
    _add_field(doc, "   Đội trưởng PCCC cơ sở",
               _v(data, "doi_truong_ten"))
    _add_field(doc, "   SĐT", _v(data, "doi_truong_sdt"))
    _add_field(doc, "   Tổng đội viên PCCC cơ sở",
               _v(data, "doi_tong_doi_vien"))

    _add_p(doc,
           "b) Nội dung và kết quả tự kiểm tra định kỳ PCCC:",
           bold=True, size=11)
    if _v(data, "lich_su_pccc"):
        _add_p(doc, _v(data, "lich_su_pccc"),
               size=11, align="justify")
    else:
        for _ in range(3):
            _add_p(doc, "…" * 80, size=11)

    _add_p(doc,
           "c) Việc bảo quản, bảo dưỡng phương tiện PCCC, CNCH:",
           bold=True, size=11)
    _add_field(doc, "   HĐ bảo dưỡng PCCC định kỳ",
               "Có" if _bool_chk(_v(data, "hd_bao_duong"))
               else "Chưa có")
    _add_field(doc, "   Đơn vị bảo dưỡng",
               _v(data, "hd_bao_duong_ncc"))

    _add_p(doc,
           "d) Việc thực hiện bảo hiểm cháy nổ bắt buộc "
           "(với cơ sở thuộc diện phải mua):",
           bold=True, size=11)
    _add_field(doc, "   Công ty bảo hiểm",
               _v(data, "bh_cong_ty"))
    _add_field(doc, "   Số HĐ bảo hiểm", _v(data, "bh_so_hd"))
    _add_field(doc, "   Ngày hết hạn",
               _v(data, "bh_ngay_het_han"))

    _add_p(doc,
           "đ) Sơ hở, thiếu sót đã phát hiện + cam kết khắc phục:",
           bold=True, size=11)
    if _v(data, "buoc_tiep_theo"):
        _add_p(doc, _v(data, "buoc_tiep_theo"),
               size=11, align="justify")
    else:
        for _ in range(3):
            _add_p(doc, "…" * 80, size=11)

    doc.add_paragraph()
    _add_section_header(doc,
                        "2. Kết quả thực hiện yêu cầu, kiến nghị "
                        "của cơ quan có thẩm quyền (nếu có):")
    for _ in range(4):
        _add_p(doc, "…" * 80, size=11)

    doc.add_paragraph()
    _add_section_header(doc,
                        "3. Đề xuất với cơ quan quản lý (nếu có):")
    for _ in range(3):
        _add_p(doc, "…" * 80, size=11)

    _add_signature(doc, vai_tro="ĐẠI DIỆN CƠ SỞ",
                   name=nguoi_dung_dau)

    doc.save(file_out)
    return file_out
