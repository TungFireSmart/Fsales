from docx import Document
from pathlib import Path

p = Path(r'D:\Fsales_PCCC\mau_hop_dong.docx')
backup = Path(r'D:\Fsales_PCCC\mau_hop_dong.backup_before_item_table.docx')
if not backup.exists():
    backup.write_bytes(p.read_bytes())

doc = Document(str(p))

anchor = None
for para in doc.paragraphs:
    t = (para.text or '').strip().lower()
    if 'điều 1' in t or 'dieu 1' in t:
        anchor = para
        break

# Create heading paragraph + table
note = doc.add_paragraph('Danh mục hàng hóa theo báo giá:')
table = doc.add_table(rows=2, cols=8)
headers = ['STT', 'Tên hàng', 'Model', 'ĐVT', 'SL', 'Đơn giá', 'Thuế', 'Thành tiền']
for i, h in enumerate(headers):
    table.cell(0, i).text = h

sample = ['{{ITEM_TABLE}}', '', '', '', '', '', '', '']
for i, v in enumerate(sample):
    table.cell(1, i).text = v

if anchor is not None:
    anchor._p.addnext(note._p)
    note._p.addnext(table._tbl)

# if not found, keep table at end (already appended)
doc.save(str(p))
print('updated_template', p)
print('backup', backup)
