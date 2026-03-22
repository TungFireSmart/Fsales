from pathlib import Path
from docx import Document

md_path = Path(r"D:\Fsales_PCCC\HUONG_DAN_CHUC_NANG_MOI_TRA_HANG_HOP_DONG.md")
out_path = Path(r"D:\Fsales_PCCC\HUONG_DAN_CHUC_NANG_MOI_TRA_HANG_HOP_DONG.docx")

text = md_path.read_text(encoding='utf-8')

doc = Document()
for line in text.splitlines():
    s = line.rstrip()
    if not s:
        doc.add_paragraph('')
        continue
    if s.startswith('# '):
        doc.add_heading(s[2:].strip(), level=1)
    elif s.startswith('## '):
        doc.add_heading(s[3:].strip(), level=2)
    elif s.startswith('### '):
        doc.add_heading(s[4:].strip(), level=3)
    elif s.startswith('- [ ] '):
        doc.add_paragraph('☐ ' + s[6:].strip(), style='List Bullet')
    elif s.startswith('- '):
        doc.add_paragraph(s[2:].strip(), style='List Bullet')
    elif s[0].isdigit() and s[1:3] == '. ':
        doc.add_paragraph(s, style='List Number')
    else:
        doc.add_paragraph(s)

doc.save(str(out_path))
print(out_path)
